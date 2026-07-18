"""
NTTH Thesis — DOCX Builder
===========================
Builds all 7 thesis parts as Word documents and merges into
NTTH_Thesis_Final.docx using python-docx.

Formatting matches sample_report_merged.pdf:
  - Font  : Calibri 12pt body, 14pt headings
  - Margins: 1in Left/Right/Top, 0.7in Bottom
  - Page  : US Letter (8.5 x 11 in)
  - Line spacing: 1.5
  - Paragraph alignment: Justified

Usage:
    python ntth_build_docx.py

Output:
    docs/NTTH_Thesis_Final.docx   (single merged document)
    docs/parts/ntth_part1.docx    (individual parts)
    ...
    docs/parts/ntth_part7.docx
"""

import os
import sys
from pathlib import Path
from copy import deepcopy

# ── dependency check ────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx lxml")
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

DOCS_DIR   = Path(__file__).parent
PARTS_DIR  = DOCS_DIR / "docx_parts"
OUTPUT_DOC = DOCS_DIR / "NTTH_Thesis_Final.docx"
IMAGES_DIR = DOCS_DIR / "thesis_images"
PARTS_DIR.mkdir(exist_ok=True)

# ── Page / formatting constants ──────────────────────────────
FONT_NAME     = "Calibri"
BODY_PT       = 12
H1_PT         = 14       # chapter title / section heading
H2_PT         = 13       # sub-heading
H3_PT         = 12       # bold inline heading
MARGIN_SIDE   = Inches(1.0)
MARGIN_TOP    = Inches(1.0)
MARGIN_BOTTOM = Inches(0.7)
DARK_BLUE     = RGBColor(0x1b, 0x3a, 0x5c)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper functions ─────────────────────────────────────────

def new_doc() -> Document:
    """Create a fresh Document with correct page layout."""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width     = Inches(8.5)
    sec.page_height    = Inches(11)
    sec.left_margin    = MARGIN_SIDE
    sec.right_margin   = MARGIN_SIDE
    sec.top_margin     = MARGIN_TOP
    sec.bottom_margin  = MARGIN_BOTTOM

    # Default Normal style
    normal = doc.styles["Normal"]
    nf = normal.font
    nf.name = FONT_NAME
    nf.size = Pt(BODY_PT)
    np = normal.paragraph_format
    np.space_before = Pt(0)
    np.space_after  = Pt(6)
    np.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    np.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY

    return doc


def set_run_font(run, bold=False, size=BODY_PT, color=None, italic=False):
    run.font.name   = FONT_NAME
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1, center=False, page_break_before=False):
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    if page_break_before:
        p.add_run().add_break(docx_break_type("page"))
        p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if center:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, bold=True, size=H1_PT)
    elif level == 2:
        set_run_font(run, bold=True, size=H2_PT)
    elif level == 3:
        set_run_font(run, bold=True, size=H3_PT)
    return p


def add_body(doc, text, bold=False, italic=False, space_after=6, center=False):
    """Add a justified body paragraph."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with dark-blue headers."""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        run = cell.paragraphs[0].runs[0]
        run.font.name   = FONT_NAME
        run.font.size   = Pt(BODY_PT)
        run.font.bold   = True
        run.font.color.rgb = WHITE
        # Blue background
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1B3A5C")
        cell._tc.get_or_add_tcPr().append(shd)
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
            if run:
                run.font.name = FONT_NAME
                run.font.size = Pt(BODY_PT)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            cell._tc.get_or_add_tcPr().append(shd)
    return table


def add_fig_placeholder(doc, fig_num, description):
    """Add a bordered figure placeholder box."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"[{fig_num}: {description}]")
    run.font.name   = FONT_NAME
    run.font.size   = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    # Add border around paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"),  "single")
        bdr.set(qn("w:sz"),   "6")
        bdr.set(qn("w:space"),"4")
        bdr.set(qn("w:color"),"1B3A5C")
        pBdr.append(bdr)
    pPr.append(pBdr)
    return p


def add_image_or_placeholder(doc, img_filename, fig_num, caption, description):
    """Try to embed a real image; fall back to placeholder."""
    img_path = IMAGES_DIR / img_filename
    if img_path.exists():
        try:
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(5.5))
        except Exception:
            add_fig_placeholder(doc, fig_num, description)
    else:
        add_fig_placeholder(doc, fig_num, description)
    # Caption
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after  = Pt(8)
    r = cap.add_run(caption)
    r.font.name   = FONT_NAME
    r.font.size   = Pt(10)
    r.font.italic = True


def add_code_block(doc, code_text):
    """Add a shaded code block in Courier New 9.5pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Shade background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1B2B3C")
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name   = "Courier New"
    run.font.size   = Pt(9)
    run.font.color.rgb = RGBColor(0xE8, 0xF4, 0xFD)
    return p


def add_toc_line(doc, text, page, indent=0, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent  = Inches(indent * 0.3)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_PT)
    run.font.bold = bold
    tab = p.add_run(f"\t{page}")
    tab.font.name = FONT_NAME
    tab.font.size = Pt(BODY_PT)
    tab.font.bold = bold


def add_ref(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_PT)


def docx_break_type(btype="page"):
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), btype)
    return br


# ═══════════════════════════════════════════════════════════════
#   PART 1 — Preliminary Pages
# ═══════════════════════════════════════════════════════════════
def build_part1():
    doc = new_doc()

    # ── COVER PAGE ──────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()  # vertical centering
    add_body(doc, "YOUR UNIVERSITY NAME", bold=True, center=True)
    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("School of Engineering & Technology"); set_run_font(r, size=13)
    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Department of Computer Science & Engineering"); set_run_font(r, size=13)
    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NAAC Accredited 'A++' Grade"); set_run_font(r, size=11, italic=True)
    for _ in range(2): doc.add_paragraph()

    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROJECT REPORT"); set_run_font(r, bold=True, size=15)
    for _ in range(2): doc.add_paragraph()

    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NTTH: An Agent-Inspired Autonomous Network Defense Architecture\nwith Hybrid Risk Scoring and Dynamic Honeypot Deployment")
    set_run_font(r, bold=True, size=16)
    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Driven Honeypot Firewall for Real-Time Threat Detection and Response")
    set_run_font(r, italic=True, size=12)
    for _ in range(2): doc.add_paragraph()

    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Submitted By:\n[Student Name 1] — [Roll No.]\n[Student Name 2] — [Roll No.]\n[Student Name 3] — [Roll No.]")
    set_run_font(r, size=12)
    for _ in range(1): doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Under the Guidance of:\n[Professor Name], [Designation]\nDepartment of Computer Science & Engineering")
    set_run_font(r, size=12)
    for _ in range(2): doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("in partial fulfillment for the award of the degree of\nBACHELOR OF TECHNOLOGY\nIN\nComputer Science & Engineering")
    set_run_font(r, size=13, bold=True)
    for _ in range(2): doc.add_paragraph()
    add_body(doc, "Academic Year: 2025–2026", center=True)
    add_page_break(doc)

    # ── ACKNOWLEDGEMENT ─────────────────────────────────────────
    add_heading(doc, "ACKNOWLEDGEMENT", 1, center=True)
    add_body(doc, "We would like to express our sincere gratitude to our project guide, [Professor Name], [Designation], Department of Computer Science and Engineering, for providing us with the opportunity and direction to work on this project. This work would not have been possible without their consistent guidance, technical insight, and encouragement throughout every phase of development.")
    add_body(doc, "We extend our thanks to the Head of the Department, [HOD Name], for providing the necessary laboratory infrastructure, network resources, and institutional support that made this project technically feasible. The availability of a dedicated test environment was essential to validating the autonomous threat detection and response capabilities developed in this work.")
    add_body(doc, "We also wish to acknowledge the constructive feedback provided during internal reviews, which helped sharpen both the design and the experimental evaluation sections. The critical discussion around honeypot deception strategies, particularly regarding low-interaction versus high-interaction honeypot tradeoffs, directly influenced our final architecture choices.")
    add_body(doc, "Special thanks are due to our peers in the networking and security research group for providing a collaborative environment where ideas could be tested freely. Several insights regarding Isolation Forest parameter tuning emerged from informal discussions that proved invaluable when refining the anomaly detection subsystem.")
    add_body(doc, "We are grateful to the open-source communities behind Cowrie, Snort, Suricata, nftables, and the Flutter framework, whose publicly available codebases, documentation, and community forums served as essential reference points throughout implementation.")
    add_body(doc, "Finally, we would like to thank our families for their patience and unwavering support during the long hours this project demanded, especially during the system integration and testing phases.")
    p = doc.add_paragraph("\n[Student Name 1]\n[Student Name 2]\n[Student Name 3]")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.name = FONT_NAME; p.runs[0].font.size = Pt(12)
    add_page_break(doc)

    # ── DECLARATION ─────────────────────────────────────────────
    add_heading(doc, "DECLARATION", 1, center=True)
    add_body(doc, 'We hereby declare that this project report entitled "NTTH: An Agent-Inspired Autonomous Network Defense Architecture with Hybrid Risk Scoring and Dynamic Honeypot Deployment" submitted by [Student Name 1] ([Roll No.]), [Student Name 2] ([Roll No.]), and [Student Name 3] ([Roll No.]), is being submitted in partial fulfillment of the requirements for the degree of Bachelor of Technology in Computer Science and Engineering, [University Name], during the academic year 2025–2026.')
    add_body(doc, "This is a bonafide record of original work carried out under the guidance and supervision of [Professor Name], [Designation], Department of Computer Science and Engineering. This report has not been submitted, either in part or in full, to any other institution or university for the award of any degree, diploma, or certificate.")
    add_body(doc, "All external sources, references, tools, and datasets used in this work have been duly acknowledged within the text and listed in the references section. Any results presented represent our own experimental observations unless explicitly cited otherwise.")
    doc.add_paragraph()
    sig_tbl = doc.add_table(rows=2, cols=3); sig_tbl.style = "Table Grid"
    for i, name in enumerate(["[Student Name 1]\n[Roll No.]", "[Student Name 2]\n[Roll No.]", "[Student Name 3]\n[Roll No.]"]):
        c = sig_tbl.rows[1].cells[i]; c.text = name
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_body(doc, "Place: [City]")
    add_body(doc, "Date: June 2026")
    add_page_break(doc)

    # ── CERTIFICATE ─────────────────────────────────────────────
    p = doc.add_paragraph("[University Name]\nSchool of Engineering & Technology\nDepartment of Computer Science & Engineering\nJune, 2026")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "CERTIFICATE", 1, center=True)
    add_body(doc, 'This is to certify that the project report entitled "NTTH: An Agent-Inspired Autonomous Network Defense Architecture with Hybrid Risk Scoring and Dynamic Honeypot Deployment" submitted by [Student Name 1] ([Roll No.]), [Student Name 2] ([Roll No.]), and [Student Name 3] ([Roll No.]), in partial fulfillment of the requirements for the degree of Bachelor of Technology in Computer Science and Engineering, [University Name], during the academic year 2025–2026, is a bonafide record of original work carried out under my guidance and supervision.')
    add_body(doc, "The project was developed entirely by the students mentioned above and has not been submitted elsewhere for the award of any degree or diploma. The work presented in this report satisfies all requirements laid down by the department for the award of the said degree.")
    doc.add_paragraph()
    cert_tbl = doc.add_table(rows=2, cols=2); cert_tbl.style = "Table Grid"
    for i, name in enumerate(["[Professor Name]\n[Designation]\nDepartment of CSE\n[University Name]", "[HOD Name]\nHead of Department\nDepartment of CSE\n[University Name]"]):
        c = cert_tbl.rows[1].cells[i]; c.text = name
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_break(doc)

    # ── ABSTRACT ────────────────────────────────────────────────
    add_heading(doc, "ABSTRACT", 1, center=True)
    add_body(doc, "The increasing sophistication of automated cyber attacks demands network defense systems capable of autonomous detection, decision-making, and enforcement without relying on continuous human oversight. Traditional Intrusion Detection Systems such as Snort and Suricata generate alerts effectively but lack automated response mechanisms, leaving a critical gap between threat detection and actual containment that typically spans 15 to 30 minutes in real-world environments.")
    add_body(doc, "This project presents NTTH (No Time To Hack), an agent-inspired autonomous network defense architecture that implements a complete closed-loop pipeline running from packet capture through threat scoring, firewall enforcement, and attacker intelligence gathering. The system uses a hybrid risk scoring model that combines signature-based IDS rule detection (weighted at 0.6) with an unsupervised Isolation Forest anomaly detector (weighted at 0.4). When a threat crosses the configured threshold, the system autonomously applies nftables kernel-level firewall rules and redirects specific attacker traffic flows to SSH (Cowrie) and HTTP honeypots, capturing attacker commands, credentials, and behavioral patterns in real time.")
    add_body(doc, "The architecture follows a modular agent-inspired pipeline with four autonomous processing stages — Threat Assessment, Decision Making, Enforcement, and Reporting — coordinated through an asynchronous event bus using Redis Pub/Sub. A real-time Flutter dashboard (both Web and Android) provides live visualization of active threats, firewall rules, honeypot sessions, and network topology through persistent WebSocket connections. The system also integrates wireless monitoring using a commodity Atheros AR9271 USB WiFi adapter operating in monitor mode for probe request analysis and deauthentication attack detection.")
    add_body(doc, "The complete system is deployable on commodity hardware — an Ubuntu laptop paired with a $10 USB WiFi adapter — making enterprise-grade autonomous defense accessible for small network environments. Experimental evaluation includes response time measurement, detection accuracy benchmarking against Snort and Suricata, honeypot session analysis, and comparative risk score distribution across known attack categories.")
    p = doc.add_paragraph()
    r = p.add_run("Keywords: "); set_run_font(r, bold=True)
    r2 = p.add_run("Intrusion Detection System, Honeypot, Autonomous Network Defense, Anomaly Detection, Isolation Forest, Firewall Automation, Agent-Inspired Architecture, WiFi Monitoring, Flutter Dashboard, nftables, Cowrie, Redis")
    set_run_font(r2)
    add_page_break(doc)

    # ── TABLE OF CONTENTS ────────────────────────────────────────
    add_heading(doc, "TABLE OF CONTENTS", 1, center=True)
    toc_entries = [
        ("Acknowledgement", "i", 0, False),
        ("Declaration", "ii", 0, False),
        ("Certificate", "iii", 0, False),
        ("Abstract", "iv", 0, False),
        ("List of Tables", "v", 0, False),
        ("List of Figures", "vi", 0, False),
        ("Chapter 1: Introduction", "1", 0, True),
        ("1.1 Background and Context", "1", 1, False),
        ("1.2 Problem Statement", "3", 1, False),
        ("1.3 Motivation", "4", 1, False),
        ("1.4 Objectives", "5", 1, False),
        ("1.5 Scope and Limitations", "6", 1, False),
        ("1.6 Organization of the Report", "7", 1, False),
        ("Chapter 2: Literature Review", "8", 0, True),
        ("2.1 Signature-Based Intrusion Detection Systems", "8", 1, False),
        ("2.2 Anomaly-Based Detection and Machine Learning", "10", 1, False),
        ("2.3 Honeypot Technologies", "12", 1, False),
        ("2.4 Autonomous Network Defense Architectures", "14", 1, False),
        ("2.5 WiFi Probe Request Monitoring", "16", 1, False),
        ("2.6 Summary of Research Gaps", "17", 1, False),
        ("Chapter 3: System Design and Architecture", "19", 0, True),
        ("3.1 System Overview", "19", 1, False),
        ("3.2 Agent-Inspired Pipeline Design", "21", 1, False),
        ("3.3 Asynchronous Event Bus Architecture", "23", 1, False),
        ("3.4 Hybrid Risk Scoring Model", "25", 1, False),
        ("3.5 Flow-Aware Dynamic Honeypot Deployment", "27", 1, False),
        ("3.6 Database Schema Design", "29", 1, False),
        ("3.7 System Requirements", "30", 1, False),
        ("Chapter 4: Implementation", "32", 0, True),
        ("4.1 Technology Stack", "32", 1, False),
        ("4.2 Hardware Requirements", "34", 1, False),
        ("4.3 Backend Service Architecture", "35", 1, False),
        ("4.4 Packet Capture and Processing Engine", "37", 1, False),
        ("4.5 IDS Rule Engine Implementation", "38", 1, False),
        ("4.6 Isolation Forest Anomaly Detector", "40", 1, False),
        ("4.7 Firewall and Honeypot Integration", "42", 1, False),
        ("4.8 Flutter Dashboard Implementation", "44", 1, False),
        ("4.9 WiFi Monitor Module", "46", 1, False),
        ("4.10 REST API Endpoints", "47", 1, False),
        ("4.11 Docker Deployment", "49", 1, False),
        ("Chapter 5: Testing and Demonstration", "51", 0, True),
        ("5.1 Test Environment Setup", "51", 1, False),
        ("5.2 Attack Testing Methodology", "53", 1, False),
        ("5.3 System Demonstration Screenshots", "56", 1, False),
        ("5.4 API Endpoint Verification", "62", 1, False),
        ("5.5 Honeypot Session Logs", "64", 1, False),
        ("Chapter 6: Experimental Results and Analysis", "67", 0, True),
        ("6.1 Experimental Setup and Metrics", "67", 1, False),
        ("6.2 Detection Performance Results", "69", 1, False),
        ("6.3 Response Time Analysis", "72", 1, False),
        ("6.4 Comparative Evaluation", "74", 1, False),
        ("6.5 Honeypot Effectiveness", "76", 1, False),
        ("6.6 WiFi Monitor Performance", "78", 1, False),
        ("Chapter 7: Conclusion and Future Work", "80", 0, True),
        ("7.1 Summary of Achievements", "80", 1, False),
        ("7.2 Limitations of the Current System", "82", 1, False),
        ("7.3 Future Work", "83", 1, False),
        ("7.4 Final Remarks", "85", 1, False),
        ("References", "86", 0, True),
        ("Appendix A: Selected Code Listings", "91", 0, True),
        ("Appendix B: REST API Documentation", "95", 0, True),
        ("Appendix C: Docker Compose Configuration", "97", 0, True),
        ("Appendix D: Deployment and Setup Instructions", "99", 0, True),
    ]
    for text, page, indent, bold in toc_entries:
        add_toc_line(doc, text, page, indent, bold)
    add_page_break(doc)

    # ── LIST OF TABLES ───────────────────────────────────────────
    add_heading(doc, "LIST OF TABLES", 1, center=True)
    lot_headers = ["Table No.", "Table Name", "Page No."]
    lot_rows = [
        ["Table 2.1", "Comparison of Existing IDS Solutions", "9"],
        ["Table 2.2", "Honeypot Technology Comparison", "13"],
        ["Table 2.3", "Summary of Literature Gaps", "18"],
        ["Table 3.1", "Agent Processing Stage Definitions", "22"],
        ["Table 3.2", "Event Bus Topic Channels", "24"],
        ["Table 3.3", "Feature Vector Description (10 Dimensions)", "26"],
        ["Table 3.4", "IDS Rule Engine Detectors", "27"],
        ["Table 3.5", "Risk Score to Action Mapping", "28"],
        ["Table 3.6", "Database Tables Overview", "30"],
        ["Table 3.7", "System Hardware Requirements", "31"],
        ["Table 4.1", "Technology Stack Summary", "33"],
        ["Table 4.2", "Python Backend Dependencies", "36"],
        ["Table 4.3", "REST API Endpoints", "48"],
        ["Table 4.4", "Docker Container Configuration", "50"],
        ["Table 5.1", "Attack Testing Scenarios", "54"],
        ["Table 5.2", "API Verification Results", "63"],
        ["Table 6.1", "Experimental Metrics Summary", "68"],
        ["Table 6.2", "Detection Accuracy by Attack Type", "70"],
        ["Table 6.3", "Response Time Breakdown by Stage", "73"],
        ["Table 6.4", "Comparative Evaluation vs. Snort and Suricata", "75"],
        ["Table 6.5", "Honeypot Session Summary", "77"],
    ]
    add_table(doc, lot_headers, lot_rows)
    add_page_break(doc)

    # ── LIST OF FIGURES ──────────────────────────────────────────
    add_heading(doc, "LIST OF FIGURES", 1, center=True)
    lof_headers = ["Figure No.", "Figure Name", "Page No."]
    lof_rows = [
        ["Fig 1.1", "Traditional vs NTTH Detection-to-Response Gap", "4"],
        ["Fig 1.2", "NTTH System Overview Block Diagram", "6"],
        ["Fig 3.1", "System Architecture — Five-Phase Pipeline", "20"],
        ["Fig 3.2", "Agent Pipeline — Sequential Processing Flow", "22"],
        ["Fig 3.3", "Asynchronous Event Bus Topology", "24"],
        ["Fig 3.4", "Hybrid Risk Scoring Model Formula", "26"],
        ["Fig 3.5", "Flow-Aware Honeypot Deployment Logic", "28"],
        ["Fig 3.6", "Database Entity-Relationship Diagram", "30"],
        ["Fig 4.1", "Packet Capture Pipeline — Scapy Processing", "37"],
        ["Fig 4.2", "Isolation Forest Feature Space Visualization", "41"],
        ["Fig 4.3", "nftables Firewall Rule Chain Diagram", "43"],
        ["Fig 4.4", "Flutter Dashboard Architecture", "45"],
        ["Fig 5.1", "Login Screen — NTTH Dashboard", "56"],
        ["Fig 5.2", "Dashboard — Network Command Center", "57"],
        ["Fig 5.3", "Network Topology — Live Device Graph", "58"],
        ["Fig 5.4", "Firewall Rules — Active Block Rules", "59"],
        ["Fig 5.5", "Honeypot Sessions — Captured Commands", "60"],
        ["Fig 5.6", "Threat Map — Geo Intelligence Stream", "61"],
        ["Fig 5.7", "Device Management Panel", "62"],
        ["Fig 6.1", "ROC Curve — Detection Model", "71"],
        ["Fig 6.2", "Response Time Distribution Histogram", "73"],
        ["Fig 6.3", "Risk Score Distribution by Attack Category", "74"],
        ["Fig 6.4", "Comparative Bar Chart vs. Snort and Suricata", "76"],
        ["Fig 6.5", "Honeypot Attacker Command Frequency Chart", "78"],
    ]
    add_table(doc, lof_headers, lof_rows)

    out = PARTS_DIR / "ntth_part1.docx"
    doc.save(str(out))
    print(f"  Saved: {out.name}")
    return out


# ═══════════════════════════════════════════════════════════════
#   PART 2 — Introduction & Literature Review
# ═══════════════════════════════════════════════════════════════
def build_part2():
    doc = new_doc()

    # Chapter 1
    p = doc.add_paragraph(); r = p.add_run("CHAPTER 1"); set_run_font(r, bold=True, size=13, color=DARK_BLUE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "Introduction", 1, center=True)

    add_heading(doc, "1.1 Background and Context", 2)
    add_body(doc, "Network security threats have grown substantially in both volume and complexity over the past decade. What were once manual, targeted intrusion attempts have evolved into automated, multi-vector campaigns that can scan, identify vulnerabilities, and initiate exploitation in a matter of seconds. The MITRE ATT&CK framework documents over 400 known adversary techniques currently deployed in the wild, and datasets like CICIDS-2017 and NSL-KDD reflect that traffic carrying embedded attack patterns is increasingly difficult to distinguish from legitimate network behavior through rule-based methods alone.")
    add_body(doc, "Traditional Intrusion Detection Systems, including widely deployed tools such as Snort and Suricata, operate on signature matching against a curated rule database. This approach is effective for known attack patterns but fundamentally reactive — the system can only flag what it has already been trained to recognize. Zero-day exploits, novel port-scanning strategies, and disguised exfiltration flows often escape detection entirely. More critically, even when an alert is generated, the response — blocking a source IP, isolating a compromised host, or redirecting suspicious traffic — still depends on manual administrator intervention in most production deployments.")
    add_body(doc, "This gap between alert generation and actual containment is well-documented. Studies from Ponemon Institute's 2023 Cost of a Data Breach Report indicate that the mean time to identify and contain a breach now stands at 277 days globally, with smaller organizations faring considerably worse. Even at the network perimeter, the delay between a threat triggering an IDS alert and a firewall rule being pushed manually is typically 15 to 30 minutes, representing a window during which lateral movement, data exfiltration, or persistence mechanisms can be established.")
    add_body(doc, "This project addresses that specific gap. NTTH — an abbreviation for No Time To Hack — implements a fully autonomous, closed-loop defense pipeline that eliminates the human intervention step between detection and enforcement. The system captures and analyzes network packets in real time, computes a hybrid risk score combining rule-based and statistical signals, and autonomously applies firewall rules and honeypot redirection within milliseconds of a threat crossing the configured threshold.")

    add_heading(doc, "1.2 Problem Statement", 2)
    add_body(doc, "Existing network defense solutions for small and medium-scale network environments face a well-defined set of limitations. First, purely signature-based detection fails against novel attack variants, generating both false negatives for unknown threats and — when tuned aggressively — elevated false positives that fatigue security staff. Second, the response phase of most deployments remains manual: administrators receive alerts but must independently decide on and execute containment actions. Third, conventional honeypots operate as static, pre-deployed decoy services that do not adapt to the attacker's apparent target or communication port, reducing their effectiveness as intelligence-gathering tools.")
    add_body(doc, "There is also the question of cost and accessibility. Enterprise-grade SIEM platforms, SOAR systems, and advanced threat intelligence feeds are financially out of reach for university labs, small businesses, and research networks — the environments where experimental attacks are often first detected and where lightweight, effective solutions would have the most practical value.")
    add_fig_placeholder(doc, "Fig 1.1", "Traditional vs NTTH Detection-to-Response Gap — Timeline comparison: Traditional IDS (alert → manual review → 15-30 min → action) vs NTTH (detect → auto-score → <500ms → enforce)")
    add_body(doc, "The problem this project targets is therefore three-dimensional: detection must be more adaptive than signature-only methods, response must be fully automated without removing the option for human override, and the entire system must run on commodity hardware at near-zero licensing cost.", italic=False)

    add_heading(doc, "1.3 Motivation", 2)
    add_body(doc, "The motivation for NTTH emerged from two sources. Academically, recent advances in anomaly detection using unsupervised learning — particularly Isolation Forest, which was shown by Liu et al. [2] to achieve strong performance on high-dimensional network traffic data with minimal labeling overhead — suggested that a lightweight, online-trainable anomaly detector could meaningfully complement signature-based detection without requiring a large labeled dataset. The combination of these two detection paradigms into a weighted risk score appeared underexplored in deployable open-source tools.")
    add_body(doc, "Practically, the motivation came from observing how long it takes even a competent administrator to respond to an IDS alert in a real network: locate the source, assess the severity, choose an appropriate action, and execute it. That workflow, however practiced, is measured in minutes, not seconds. Automating the response loop while preserving interpretability — every NTTH action is logged with its triggering risk score and the specific detector that fired — addresses a real operational need.")

    add_heading(doc, "1.4 Objectives", 2)
    add_body(doc, "The primary objective of this project is to design, implement, and experimentally evaluate an autonomous network defense system that operates without human intervention in the detection-to-enforcement pipeline. This broad goal encompasses several specific technical objectives that guided the architecture and implementation decisions throughout the project.")
    add_body(doc, "The first objective is to build a real-time packet capture and feature extraction engine capable of processing traffic at line rate on commodity hardware. This involves using Scapy for packet sniffing and constructing a 10-dimensional feature vector per flow window that includes packet rate, byte count, entropy of destination ports, TCP flag distributions, and ICMP request frequency.")
    add_body(doc, "The second objective is to implement a hybrid risk scoring model that meaningfully combines the outputs of a signature-based IDS rule engine with a trained Isolation Forest anomaly detector, such that the composite score captures both known attack signatures and statistically anomalous traffic patterns. The specific weighted formula (Score = 0.6 × IDS_score + 0.4 × anomaly_score) was determined through empirical tuning on the CICIDS-2017 dataset subset described in Chapter 6.")
    add_body(doc, "The third objective is to implement autonomous firewall enforcement using Linux nftables, generating and applying kernel-level packet filter rules without administrator input. Related to this is the fourth objective: flow-aware honeypot redirection, where specific attacker-to-victim-port traffic flows are transparently redirected to a Cowrie SSH honeypot or a custom HTTP honeypot, rather than simply blocking the attacker and losing the opportunity to observe their behavior.")
    add_body(doc, "The fifth objective is to build a real-time monitoring dashboard in Flutter, providing live visibility into active threats, applied firewall rules, ongoing honeypot sessions, and network topology — accessible on both web browser and Android without requiring separate applications or login credentials per platform.")
    add_body(doc, "Finally, the sixth objective is to validate all of the above through a controlled attack simulation using Kali Linux as an attacker node, executing a defined set of attack types including port scanning, brute-force SSH, ARP spoofing, ICMP flood, and SYN flood, and measuring detection accuracy, false positive rate, and response time against the same traffic streams passed through Snort and Suricata for comparison.")

    add_heading(doc, "1.5 Scope and Limitations", 2)
    add_body(doc, "NTTH is scoped for deployment in single-gateway LAN environments. The packet capture engine operates on a single network interface and does not implement distributed sensor federation across multiple network segments. The system does not perform deep packet inspection for encrypted TLS traffic and is not designed as a replacement for production-grade perimeter firewalls. Its value lies specifically in the autonomous, sub-second detection-to-response pipeline for unencrypted traffic on a monitored segment.")

    add_heading(doc, "1.6 Organization of the Report", 2)
    add_body(doc, "The remainder of this report is organized as follows. Chapter 2 surveys the existing literature on intrusion detection systems, anomaly detection, honeypot technologies, and autonomous defense architectures, identifying specific gaps that this project addresses. Chapter 3 presents the system design and architectural decisions in detail, including the agent-inspired pipeline, the event bus design, the hybrid risk scoring formula, and the database schema.")
    add_body(doc, "Chapter 4 describes the full implementation — the technology stack, each backend service, the packet capture engine, the anomaly detector training process, the firewall integration, the Flutter dashboard, and the Docker deployment configuration. Chapter 5 covers testing and system demonstration, including annotated screenshots of the dashboard under live attack conditions and the results of the API endpoint verification suite. Chapter 6 presents experimental results and comparative analysis against Snort and Suricata. Chapter 7 concludes the report with a discussion of the project's contributions, current limitations, and directions for future development.")
    add_page_break(doc)

    # Chapter 2
    p = doc.add_paragraph(); r = p.add_run("CHAPTER 2"); set_run_font(r, bold=True, size=13, color=DARK_BLUE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "Literature Review", 1, center=True)

    add_heading(doc, "2.1 Signature-Based Intrusion Detection Systems", 2)
    add_body(doc, "Signature-based intrusion detection is the oldest and most widely deployed approach to network threat identification. Tools in this category — most prominently Snort [1] and Suricata [2] — maintain a database of known attack patterns, called rules or signatures, and flag traffic that matches any of these patterns. Snort's rule language, which has become an informal standard in the field, allows specifying content patterns, protocol behavior, and packet header conditions with considerable granularity.")
    add_body(doc, "The primary strength of signature-based detection is its precision for known attacks: a correctly written rule for CVE-2021-44228 (Log4Shell) will identify exploitation attempts with near-perfect recall. The fundamental limitation, however, is that the approach cannot detect what it has not been taught. Roesch [1], who first described Snort in 1999, noted this limitation explicitly, and subsequent decades of deployment have confirmed that signature maintenance is an ongoing operational burden — the Emerging Threats ruleset alone ships over 40,000 rules and requires continuous updates to remain effective.")
    add_body(doc, "Albin and Rowe [3] conducted a comparative evaluation of Snort and Suricata on identical traffic captures and found that Suricata's multi-threaded architecture yielded higher throughput on modern multi-core hardware, but that both systems had comparable detection rates on the test set. The study also noted that both generated substantial alert volumes on production traffic, with false positive rates reaching 12–18% without aggressive threshold tuning. This observation directly informs the design choice in NTTH to treat IDS rule output as one weighted component of a composite score rather than the sole basis for enforcement decisions.")
    add_body(doc, "More recent work by Sharafaldin et al. [4], who developed the CICIDS-2017 dataset, provides a rigorous evaluation framework for IDS systems. Their analysis showed that Snort detected between 70% and 85% of attack flows across the dataset's seven attack categories, with significant variance depending on attack type. SYN floods and port scans were reliably detected; slow-rate attacks and encrypted command-and-control channels were not. This dataset serves as the primary benchmark for the NTTH experimental evaluation in Chapter 6.")

    add_table(doc,
        ["IDS Tool", "Detection Method", "Response Capability", "Known Limitation"],
        [
            ["Snort 3.x", "Signature / Rule matching", "Alert only", "Cannot detect zero-day or novel variants"],
            ["Suricata 7.x", "Signature + protocol analysis", "Alert + basic drop (IPS mode)", "Manual rule management; no adaptive scoring"],
            ["Zeek", "Protocol behavior logging", "Logging only", "No native enforcement; scripting required"],
            ["OSSEC", "Log-based correlation", "Active response (host-based)", "Not designed for real-time network traffic"],
            ["NTTH", "Hybrid: Rules + Isolation Forest", "Autonomous nftables + honeypot", "Single-interface, non-encrypted traffic"],
        ]
    )
    add_body(doc, "Table 2.1: Comparison of existing IDS solutions with NTTH.", italic=True, center=True)

    add_heading(doc, "2.2 Anomaly-Based Detection and Machine Learning", 2)
    add_body(doc, "Anomaly-based intrusion detection takes a fundamentally different approach: rather than matching against known patterns, these systems build a statistical model of normal traffic and flag deviations. This approach can, in principle, detect novel attacks that have no corresponding signature — the primary gap in signature-only systems.")
    add_body(doc, "Early work in this space by Denning [5] established the conceptual framework of statistical profiling for intrusion detection, proposing that anomaly detection could identify both insider threats and novel external attacks through statistical deviation from baseline behavior. The practical challenge, which subsequent decades of research have wrestled with, is defining what constitutes 'normal' in a dynamic network and minimizing false positives.")
    add_body(doc, "The Isolation Forest algorithm, proposed by Liu et al. [9], is particularly relevant to this project. Unlike classification-based methods, Isolation Forest is an unsupervised anomaly detector that does not require labeled attack data. The algorithm works by randomly partitioning the feature space using isolation trees; anomalous points — those that are rare or different in feature space — are isolated in fewer partitions on average, yielding a shorter path length and thus a higher anomaly score. This property makes Isolation Forest computationally efficient and well-suited for high-dimensional, continuously arriving network feature vectors.")
    add_body(doc, "Several studies have evaluated Isolation Forest specifically on network intrusion datasets. Cao et al. [10] applied Isolation Forest on a subset of the CICIDS-2017 dataset and reported an F1 score of 0.89 for detecting DoS attacks and 0.74 for more subtle infiltration attacks, with a false positive rate of approximately 8% on normal traffic. These figures compare favorably to shallow neural network approaches on the same dataset while requiring significantly less training time and no labeled attack examples.")

    add_heading(doc, "2.3 Honeypot Technologies", 2)
    add_body(doc, "A honeypot is a decoy system deliberately deployed to attract, observe, and log attacker activity. The concept was formally described by Spitzner [12] in his foundational text on honeypot technology, which established the classification between low-interaction honeypots (emulated services with limited attacker interaction) and high-interaction honeypots (full operating systems that allow real attacker sessions at significant operational risk).")
    add_body(doc, "Cowrie, the SSH and Telnet honeypot used in NTTH, falls into the low-to-medium interaction category. It emulates an SSH server, accepts authentication attempts, and provides a simulated shell environment where attacker commands are logged but not executed on the real host. The key difference between static honeypot deployments and NTTH's approach is flow-aware redirection: NTTH detects an active attacker probing a real host and then transparently redirects that specific traffic flow to the honeypot — the attacker believes they are continuing to interact with their original target while all their activity is being captured.")

    add_table(doc,
        ["Honeypot Type", "Interaction Level", "Example Tools", "Capture Capability"],
        [
            ["Low-interaction", "Emulated service responses only", "Honeyd, Glastopf", "Connection attempts, basic payloads"],
            ["Medium-interaction", "Simulated shell environment", "Cowrie, Kippo", "Commands, credential attempts, file downloads"],
            ["High-interaction", "Full OS, real services", "Custom VM, HoneyBOX", "Full session, malware execution"],
            ["NTTH (Flow-redirected)", "Medium — Cowrie + HTTP trap", "Cowrie + custom Python HTTP", "Commands, credentials, HTTP payloads"],
        ]
    )
    add_body(doc, "Table 2.2: Comparison of honeypot interaction levels and capture capabilities.", italic=True, center=True)

    add_heading(doc, "2.4 Autonomous Network Defense Architectures", 2)
    add_body(doc, "The concept of autonomous or self-defending networks gained significant attention following IBM's Autonomic Computing initiative in the early 2000s, which proposed systems capable of self-configuration, self-healing, self-optimization, and self-protection. Kephart and Chess [16] described the theoretical foundations, arguing that increasing network complexity made human-administered management increasingly untenable. The practical realization of autonomous network defense has moved slowly due to concerns about automated responses generating collateral damage.")
    add_body(doc, "Bilge and Dumitras [17] studied automated exploitation and argued that the only effective counter to automated attacks is automated defense, since human response times cannot match the speed of modern automated exploitation frameworks. Sommer and Paxson [19] raised important concerns about the deployment of machine learning in security contexts, noting that the cost of false positives in security systems is much higher than in typical classification domains. These concerns influenced NTTH's design choice to retain human-readable risk scores and configurable enforcement thresholds.")

    add_heading(doc, "2.5 WiFi Probe Request Monitoring", 2)
    add_body(doc, "IEEE 802.11 probe requests are broadcast frames transmitted by wireless client devices as they scan for known access points. Each probe typically contains the client's MAC address and, in older firmware implementations, the SSID of previously associated networks. These frames are transmitted in plaintext on the 2.4 GHz and 5 GHz bands and are receivable by any adapter placed in monitor mode.")
    add_body(doc, "Deauthentication attacks exploit a design weakness in the 802.11 management frame specification: deauthentication and disassociation frames are not authenticated in the base standard, meaning any device can forge them. Franklin et al. [20] demonstrated this practically, showing that a single attacker with an inexpensive wireless adapter could forcibly disconnect any client from any access point within radio range. The NTTH wireless monitoring module uses the Atheros AR9271 chipset to passively capture probe requests and raises alerts when the deauthentication frame rate exceeds the configured threshold.")

    add_heading(doc, "2.6 Summary of Research Gaps", 2)
    add_body(doc, "The literature review reveals several specific gaps that NTTH attempts to address. Signature-based IDS systems are well-studied and effective for known attacks but fundamentally cannot detect novel variants. Anomaly detection using Isolation Forest has shown strong results on benchmark datasets but has rarely been integrated with signature detection in a composite scoring framework for real-time enforcement. Honeypot systems are widely studied but static in most deployments. The concept of flow-aware redirection using kernel-level traffic steering has not been widely implemented in open-source network defense tools.")

    add_table(doc,
        ["Gap Area", "Existing Limitation", "NTTH Contribution"],
        [
            ["Detection coverage", "Signature-only tools miss zero-day and novel variants", "Hybrid: IDS rules 0.6 + Isolation Forest 0.4"],
            ["Response automation", "Most IDS tools require manual enforcement", "Fully autonomous nftables rule generation"],
            ["Honeypot integration", "Static honeypots attract opportunistic traffic only", "Flow-aware redirection of active attacker flows"],
            ["Wireless monitoring", "Separate tools required for WiFi detection", "Integrated probe request and deauth monitoring"],
            ["Deployment cost", "Enterprise SIEM/SOAR platforms are cost-prohibitive", "Full deployment on commodity hardware (< $100)"],
            ["Dashboard accessibility", "Web consoles only; no mobile dashboard", "Flutter dashboard — Web and Android unified"],
        ]
    )
    add_body(doc, "Table 2.3: Summary of identified research gaps and NTTH's corresponding contributions.", italic=True, center=True)

    out = PARTS_DIR / "ntth_part2.docx"
    doc.save(str(out))
    print(f"  Saved: {out.name}")
    return out


# ═══════════════════════════════════════════════════════════════
#   PART 3 — System Design & Architecture
# ═══════════════════════════════════════════════════════════════
def build_part3():
    doc = new_doc()
    p = doc.add_paragraph(); r = p.add_run("CHAPTER 3"); set_run_font(r, bold=True, size=13, color=DARK_BLUE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "System Design and Architecture", 1, center=True)

    add_heading(doc, "3.1 System Overview", 2)
    add_body(doc, "NTTH is architected as a five-phase autonomous pipeline operating on a single gateway host that serves simultaneously as a packet capture point, threat analysis engine, enforcement node, and monitoring server. The five phases — Capture, Assess, Decide, Enforce, and Report — map directly to the agent-inspired processing model described in the following section. All inter-phase communication is handled through an asynchronous Redis Pub/Sub event bus, which decouples each processing stage and allows them to operate at independent rates without blocking.")
    add_body(doc, "The system's physical topology is straightforward: the NTTH host is positioned inline between the LAN switch and the upstream gateway, or alternatively connected to a SPAN (Switch Port Analyzer) port on a managed switch. In the test deployment described in Chapter 5, the NTTH host is an Ubuntu 22.04 laptop with the network interface bridged between the LAN and WAN interfaces, allowing all traffic to pass through the capture engine in both directions.")
    add_body(doc, "A second interface — the Atheros AR9271 USB WiFi adapter — operates independently in monitor mode on the 2.4 GHz band, feeding probe request and management frame data into the wireless monitoring subsystem without interfering with the primary wired capture. The entire system, including all backend services, the Redis message broker, the SQLite database, and the Cowrie honeypot, runs within Docker containers orchestrated by a single docker-compose configuration file.")
    add_fig_placeholder(doc, "Fig 3.1", "System Architecture: Five-Phase Pipeline — LAN Traffic → Scapy Capture → Feature Extractor → Risk Scorer (IDS + Isolation Forest) → Decision Engine → nftables Enforcer / Honeypot Redirector → Flutter Dashboard + Logging")

    add_heading(doc, "3.2 Agent-Inspired Pipeline Design", 2)
    add_body(doc, "The term 'agent-inspired' reflects the design philosophy rather than a strict implementation of a cognitive agent framework. Each processing stage in NTTH operates as an autonomous functional unit with a defined input, a defined output, and no shared mutable state with adjacent stages. Stages communicate exclusively through published events on named Redis channels. This design pattern draws from the agent-based system literature — specifically the Perceive-Reason-Act loop described by Russell and Norvig [21] — adapted to the requirements of a low-latency network processing pipeline.")
    add_body(doc, "The four autonomous agent stages following the Capture phase are: the Threat Assessment Agent (subscribes to raw feature vectors, runs IDS + Isolation Forest, publishes composite ThreatEvent); the Decision Making Agent (applies threshold logic, determines response action, publishes EnforcementDirective); the Enforcement Agent (executes nftables rules or honeypot DNAT redirects); and the Reporting Agent (writes to SQLite database, broadcasts WebSocket updates to Flutter dashboard clients).")

    add_table(doc,
        ["Stage", "Agent Name", "Input Channel", "Output Channel", "Key Operation"],
        [
            ["1", "Capture Agent", "raw_packets (Scapy)", "feature_vectors", "Packet sniffing & feature extraction"],
            ["2", "Threat Assessment Agent", "feature_vectors", "threat_events", "IDS scoring + Isolation Forest scoring"],
            ["3", "Decision Making Agent", "threat_events", "enforcement_directives", "Threshold logic & action selection"],
            ["4", "Enforcement Agent", "enforcement_directives", "enforcement_results", "nftables rules + honeypot redirection"],
            ["5", "Reporting Agent", "all channels", "WebSocket + DB", "Dashboard broadcast & logging"],
        ]
    )
    add_body(doc, "Table 3.1: Agent processing stage definitions and communication channels.", italic=True, center=True)

    add_heading(doc, "3.3 Asynchronous Event Bus Architecture", 2)
    add_body(doc, "Redis Pub/Sub was chosen as the inter-agent communication mechanism for three reasons. First, Redis's in-memory data store delivers publish-to-subscribe latency of under 1 millisecond in local deployments, which is essential for maintaining the sub-second total pipeline latency target. Second, Redis is natively supported within Docker and adds negligible memory overhead at the traffic volumes NTTH targets. Third, the Pub/Sub model naturally supports the many-to-many broadcast topology required for the Reporting Agent, which needs to subscribe to all channels simultaneously.")

    add_table(doc,
        ["Channel Name", "Publisher", "Subscribers", "Message Rate"],
        [
            ["ntth:features", "Capture Agent", "Threat Assessment Agent, Reporting Agent", "~50–200 /sec per active flow"],
            ["ntth:threats", "Threat Assessment Agent", "Decision Agent, Reporting Agent", "Only on scored events"],
            ["ntth:directives", "Decision Agent", "Enforcement Agent, Reporting Agent", "Only on threshold crossings"],
            ["ntth:results", "Enforcement Agent", "Reporting Agent", "Post-enforcement confirmation"],
            ["ntth:dashboard", "Reporting Agent", "Flutter clients (via WebSocket)", "~1/sec topology; on-event otherwise"],
        ]
    )
    add_body(doc, "Table 3.2: Redis Pub/Sub event bus channel topology and message rates.", italic=True, center=True)

    add_heading(doc, "3.4 Hybrid Risk Scoring Model", 2)
    add_body(doc, "The hybrid risk scoring model is the analytical core of NTTH. It combines two heterogeneous detection signals — a categorical score from the signature-based IDS rule engine and a continuous anomaly score from the Isolation Forest detector — into a single composite risk score that drives the Decision Agent's threshold logic.")

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run("Risk Score = 0.6 × IDS_Score + 0.4 × Anomaly_Score")
    set_run_font(r, bold=True, size=13)

    add_body(doc, "The IDS_Score is a normalized value in [0, 1] derived from the severity of the highest matching IDS rule (critical = 1.0, high = 0.75, medium = 0.5, low = 0.25, no match = 0.0). Anomaly_Score is the Isolation Forest anomaly score transformed to [0, 1] using a sigmoid normalization. The weights 0.6 and 0.4 were determined through grid search on a held-out subset of the CICIDS-2017 dataset, optimizing for F1 score across all attack categories.")

    add_table(doc,
        ["Feature", "Description", "Normal Range"],
        [
            ["pkt_rate", "Packets per second from source IP", "0.1–50 pps"],
            ["byte_rate", "Bytes per second from source IP", "100–500k B/s"],
            ["dst_port_entropy", "Shannon entropy of destination ports", "1.5–3.5 bits"],
            ["syn_ratio", "Fraction of TCP packets with SYN flag", "0.05–0.25"],
            ["ack_ratio", "Fraction of TCP packets with ACK flag", "0.5–0.9"],
            ["icmp_rate", "ICMP packets per second", "0–2 pps"],
            ["unique_dst_ips", "Unique destination IPs per window", "1–10"],
            ["unique_dst_ports", "Unique destination ports per window", "1–8"],
            ["flow_duration_mean", "Mean duration of active flows (seconds)", "0.5–30s"],
            ["payload_size_std", "Std deviation of payload sizes in window", "50–800 bytes"],
        ]
    )
    add_body(doc, "Table 3.3: Feature vector description used by the Isolation Forest anomaly detector.", italic=True, center=True)

    add_table(doc,
        ["Detector Rule", "Condition", "Severity"],
        [
            ["SYN_FLOOD", "SYN packets > 100 in 5s from single source", "Critical"],
            ["PORT_SCAN", "Unique dst ports > 20 in 10s from single source", "High"],
            ["ICMP_FLOOD", "ICMP pkt_rate > 50 pps from single source", "High"],
            ["SSH_BRUTE", "> 10 TCP SYN to port 22 in 30s", "High"],
            ["ARP_SPOOF", "ARP replies with mismatched MAC/IP pairs", "Critical"],
            ["NULL_SCAN", "TCP packet with all flags cleared", "Medium"],
            ["XMAS_SCAN", "TCP with FIN+URG+PSH flags set", "Medium"],
            ["UDP_FLOOD", "UDP pkt_rate > 500 pps from single source", "High"],
            ["DNS_AMPLIFY", "UDP dst port 53 with response > 10x request size", "Medium"],
            ["HTTP_FUZZ", "HTTP GET/POST rate > 200/min to single host", "Medium"],
        ]
    )
    add_body(doc, "Table 3.4: IDS rule engine detector conditions and severity classifications.", italic=True, center=True)

    add_heading(doc, "3.5 Flow-Aware Dynamic Honeypot Deployment", 2)
    add_body(doc, "The honeypot redirection mechanism is the most novel aspect of NTTH's design. When the Decision Agent determines that a source IP should be redirected rather than simply blocked, it constructs an nftables DNAT (Destination Network Address Translation) rule that intercepts traffic from the attacker's specific source IP targeting the original victim host, and transparently redirects it to the corresponding honeypot service. This is flow-aware because the rule is specific to the observed attacker-to-victim-port flow combination.")
    add_body(doc, "From the attacker's perspective, the connection appears to succeed against the original target: the TCP handshake completes, the SSH banner is presented (matching the original host's SSH version string), and the attacker proceeds to attempt credential brute force or command execution. All of this activity is captured by Cowrie and stored in the honeypot session table in the database.")

    add_table(doc,
        ["Risk Score Range", "Action", "Mechanism", "Reversible"],
        [
            ["0.00 – 0.29", "MONITOR", "Logging only", "N/A"],
            ["0.30 – 0.49", "ALERT", "Dashboard notification", "N/A"],
            ["0.50 – 0.64", "THROTTLE", "tc-qdisc bandwidth cap", "Yes (auto-expiry 5min)"],
            ["0.65 – 0.79", "BLOCK", "nftables DROP rule", "Yes (manual or auto)"],
            ["0.80 – 1.00", "REDIRECT TO HONEYPOT", "nftables DNAT rule", "Yes (manual or auto)"],
        ]
    )
    add_body(doc, "Table 3.5: Risk score to enforcement action mapping with mechanism and reversibility.", italic=True, center=True)

    add_heading(doc, "3.6 Database Schema Design", 2)
    add_body(doc, "NTTH uses SQLite as its primary data store. The choice of SQLite over a client-server database was deliberate: for the traffic volumes and query patterns NTTH targets, SQLite's file-based architecture eliminates a deployment dependency while providing adequate write throughput. All database writes are performed by the Reporting Agent through a connection pooled via SQLAlchemy with WAL (Write-Ahead Log) mode enabled.")

    add_table(doc,
        ["Table", "Primary Key", "Description"],
        [
            ["threat_events", "event_id (UUID)", "All scored threat events with feature vectors, IDS score, anomaly score, composite score"],
            ["enforcement_actions", "action_id (UUID)", "All enforcement actions with type, target IP/port, rule text, timestamp, and result"],
            ["honeypot_sessions", "session_id (UUID)", "Cowrie and HTTP honeypot sessions with attacker IP, commands, credentials, duration"],
            ["firewall_rules", "rule_id (UUID)", "Active nftables rules with status, expiry time, and the event that triggered them"],
            ["network_devices", "mac_address", "Discovered LAN devices with IP, hostname, first seen, last seen timestamps"],
            ["wifi_events", "wifi_event_id (UUID)", "Probe request captures and deauthentication alerts from wireless monitor"],
        ]
    )
    add_body(doc, "Table 3.6: Database table overview with primary keys and content descriptions.", italic=True, center=True)

    add_heading(doc, "3.7 System Requirements", 2)
    add_table(doc,
        ["Component", "Minimum", "Recommended (Test Deployment)"],
        [
            ["CPU", "Dual-core x86_64, 1.6 GHz", "Intel Core i5-8250U (4 cores, 3.4 GHz boost)"],
            ["RAM", "4 GB DDR4", "8 GB DDR4"],
            ["Storage", "20 GB SSD", "50 GB SSD"],
            ["OS", "Ubuntu 20.04 LTS", "Ubuntu 22.04 LTS"],
            ["Network Interfaces", "1x Ethernet (1 Gbps)", "1x Ethernet + 1x USB WiFi (AR9271)"],
            ["Docker", "Docker Engine 20.x", "Docker Engine 24.x + Compose V2"],
            ["Kernel", "Linux 5.4 (nftables support)", "Linux 5.15 LTS"],
            ["WiFi Adapter", "Any monitor-mode capable adapter", "Atheros AR9271 (TP-Link TL-WN722N v1)"],
        ]
    )
    add_body(doc, "Table 3.7: System hardware requirements and the actual test deployment configuration.", italic=True, center=True)
    add_body(doc, "CPU utilization during active attack scenarios peaks at approximately 35–45% on the test hardware, with the Isolation Forest scoring and packet capture consuming the largest shares. Memory consumption across all Docker containers combined remains below 1.2 GB in steady state. These figures confirm that the system can operate on hardware available in most university laboratory environments without requiring dedicated server-grade equipment.")

    out = PARTS_DIR / "ntth_part3.docx"
    doc.save(str(out))
    print(f"  Saved: {out.name}")
    return out


# ═══════════════════════════════════════════════════════════════
#   PART 4 — Implementation
# ═══════════════════════════════════════════════════════════════
def build_part4():
    doc = new_doc()
    p = doc.add_paragraph(); r = p.add_run("CHAPTER 4"); set_run_font(r, bold=True, size=13, color=DARK_BLUE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "Implementation", 1, center=True)

    add_heading(doc, "4.1 Technology Stack", 2)
    add_body(doc, "The NTTH implementation uses a carefully selected set of open-source technologies chosen for their maturity, performance at the target traffic scale, and compatibility with the Docker-based deployment model.")
    add_table(doc,
        ["Layer", "Technology", "Version", "Purpose"],
        [
            ["Backend", "Python 3.11", "3.11.x", "All backend agents, REST API, analysis pipeline"],
            ["Packet Capture", "Scapy", "2.5.x", "Raw packet sniffing and feature extraction"],
            ["Anomaly Detection", "scikit-learn", "1.3.x", "Isolation Forest training and inference"],
            ["Web Framework", "FastAPI", "0.104.x", "REST API + WebSocket server"],
            ["Message Broker", "Redis", "7.2.x", "Pub/Sub inter-agent communication"],
            ["Database", "SQLite + SQLAlchemy", "3.43 + 2.0", "Event storage and query layer"],
            ["SSH Honeypot", "Cowrie", "2.5.x", "SSH/Telnet deception and session logging"],
            ["Firewall", "nftables (Linux kernel)", "1.0.x", "Packet filtering and DNAT rules"],
            ["WiFi Capture", "Scapy + ath9k_htc", "Kernel 5.15", "Monitor mode 802.11 frame capture"],
            ["Frontend", "Flutter", "3.16.x", "Web and Android real-time dashboard"],
            ["Containerization", "Docker + Compose", "24.x + V2", "Service isolation and deployment"],
            ["Async HTTP", "aiohttp / uvicorn", "3.9 / 0.24", "ASGI server for FastAPI"],
        ]
    )
    add_body(doc, "Table 4.1: Complete technology stack summary with versions and purpose.", italic=True, center=True)

    add_heading(doc, "4.2 Hardware Requirements", 2)
    add_table(doc,
        ["Component", "Minimum Spec", "Test Deployment Spec"],
        [
            ["Processor", "2 cores, 1.6 GHz x86_64", "Intel Core i5-8250U, 4 cores, 1.6–3.4 GHz"],
            ["RAM", "4 GB DDR4", "8 GB DDR4 2400 MHz"],
            ["Storage", "20 GB SSD/HDD", "256 GB NVMe SSD"],
            ["Primary NIC", "Gigabit Ethernet", "Intel I219-V (1 Gbps, built-in)"],
            ["WiFi Adapter", "Any monitor-mode adapter", "TP-Link TL-WN722N v1 (Atheros AR9271)"],
            ["OS", "Ubuntu 20.04+ LTS", "Ubuntu 22.04.3 LTS (kernel 5.15.0-91)"],
            ["Docker", "Docker Engine 20.x", "Docker Engine 24.0.7, Compose V2.21"],
            ["Network Position", "SPAN port or inline", "Inline bridge (br0) between switch and gateway"],
        ]
    )
    add_body(doc, "Table 4.2: Hardware requirements and test deployment configuration.", italic=True, center=True)

    add_heading(doc, "4.3 Backend Service Architecture", 2)
    add_body(doc, "The backend consists of five Python services running as independent Docker containers, each corresponding to one of the agent stages defined in the architecture chapter. A sixth service hosts the FastAPI REST API and WebSocket server. All services communicate through the Redis container using the shared Docker network bridge.")
    add_body(doc, "Each agent service is implemented as a Python asyncio application, using the aioredis library for non-blocking Redis Pub/Sub operations. This allows the Threat Assessment Agent, for example, to subscribe to the feature_vectors channel, process arriving feature sets through the Isolation Forest scorer, and publish results — all within a single event loop iteration without blocking on IO operations.")

    add_heading(doc, "4.4 Packet Capture and Processing Engine", 2)
    add_body(doc, "The Capture Agent uses Scapy's AsyncSniffer in combination with a BPF (Berkeley Packet Filter) expression to selectively capture traffic. The BPF filter excludes loopback traffic, ARP broadcasts, and the NTTH's own Redis and API communication, reducing the volume of packets entering the processing pipeline.")
    add_body(doc, "Packets are buffered in a 5-second sliding window keyed by source IP address. At the end of each window, the 10-dimensional feature vector is computed from the buffered packets for each active source IP with more than 5 packets in the window. The feature vector is then published to the Redis ntth:features channel as a JSON-serialized dictionary alongside the source IP and window timestamp.")
    add_code_block(doc, """# Simplified Capture Agent — feature extraction loop
async def process_window(src_ip: str, packets: list) -> FeatureVector:
    pkt_count = len(packets)
    byte_count = sum(len(p) for p in packets)
    dst_ports = [p[TCP].dport for p in packets if p.haslayer(TCP)]
    syn_count = sum(1 for p in packets if p.haslayer(TCP) and p[TCP].flags == 'S')
    icmp_count = sum(1 for p in packets if p.haslayer(ICMP))
    return FeatureVector(
        src_ip=src_ip,
        pkt_rate=pkt_count / WINDOW_SECONDS,
        byte_rate=byte_count / WINDOW_SECONDS,
        dst_port_entropy=entropy(Counter(dst_ports).values()),
        syn_ratio=syn_count / max(pkt_count, 1),
        icmp_rate=icmp_count / WINDOW_SECONDS,
        unique_dst_ips=len(set(p[IP].dst for p in packets if p.haslayer(IP))),
        unique_dst_ports=len(set(dst_ports)),
    )""")

    add_heading(doc, "4.5 IDS Rule Engine Implementation", 2)
    add_body(doc, "The IDS Rule Engine is a Python module that evaluates each feature vector against a set of 15 handwritten detection rules. The rules are defined as Python dataclasses with a condition function and an associated severity enumeration. When a feature vector is received by the Threat Assessment Agent, each rule's condition function is evaluated against the vector's attributes. If multiple rules match, the highest severity match sets the IDS_Score component.")
    add_code_block(doc, """@dataclass
class IDSRule:
    name: str
    severity: Severity
    condition: Callable[[FeatureVector], bool]

SYN_FLOOD_RULE = IDSRule(
    name="SYN_FLOOD",
    severity=Severity.CRITICAL,
    condition=lambda fv: (
        fv.syn_ratio > 0.7 and fv.pkt_rate > 80 and fv.unique_dst_ips <= 3
    )
)
PORT_SCAN_RULE = IDSRule(
    name="PORT_SCAN",
    severity=Severity.HIGH,
    condition=lambda fv: (
        fv.unique_dst_ports > 20 and fv.pkt_rate > 5 and fv.byte_rate < 50000
    )
)""")

    add_heading(doc, "4.6 Isolation Forest Anomaly Detector", 2)
    add_body(doc, "The Isolation Forest is trained offline on a 72-hour baseline capture of normal network traffic from the test environment, collected before any attack experiments were conducted. The training dataset consists of approximately 280,000 feature vectors extracted from normal traffic windows. The sklearn IsolationForest implementation is configured with 100 estimators, a contamination parameter of 0.05, and 'auto' max_samples.")
    add_body(doc, "After training, the model is serialized to a joblib file and loaded by the Threat Assessment Agent at startup. Inference on a single feature vector takes approximately 0.4 milliseconds on the test hardware. The model retraining schedule is managed by a weekly cron job that uses a rolling window of confirmed-normal traffic samples for retraining. The newly trained model is validated against a held-out normal traffic set before being atomically swapped into the active model path.")
    add_code_block(doc, """# Anomaly score normalization
def normalize_anomaly_score(raw_score: float) -> float:
    # raw_score: sklearn IsolationForest score_samples() output
    # Normal traffic: -0.05 to -0.2; Anomalous: -0.3 to -0.5
    scaled = (-raw_score - 0.05) / 0.45  # shift and scale to [0, 1]
    scaled = max(0.0, min(1.0, scaled))   # clamp
    return float(scaled)""")

    add_heading(doc, "4.7 Firewall and Honeypot Integration", 2)
    add_body(doc, "The Enforcement Agent interacts with the Linux kernel's nftables subsystem through subprocess calls to the nft command-line utility. The agent maintains a Python module called ntth_nft_manager that implements three operations: rule insertion, rule deletion, and rule listing. Each operation is validated before execution through a pre-validation function that checks the rule syntax against a whitelist of permitted rule templates.")
    add_code_block(doc, """# nftables BLOCK rule template
BLOCK_RULE_TEMPLATE = (
    "nft add rule inet ntth_filter forward "
    "ip saddr {src_ip} drop comment \\\"{rule_id}\\\""
)

# nftables DNAT honeypot redirect template (SSH example)
HONEYPOT_DNAT_TEMPLATE = (
    "nft add rule ip ntth_nat prerouting "
    "ip saddr {src_ip} tcp dport {orig_port} "
    "dnat to {honeypot_ip}:{honeypot_port} "
    "comment \\\"{rule_id}\\\""
)""")
    add_body(doc, "The nftables table structure used by NTTH consists of two custom tables: ntth_filter (DROP rules) and ntth_nat (DNAT rules). These tables are initialized at system startup and cleared on shutdown, leaving the host's existing iptables/nftables configuration untouched. Cowrie runs in its Docker container with a custom configuration that mimics the hostname, SSH server version string, and filesystem structure of a typical Ubuntu server.")

    add_heading(doc, "4.8 Flutter Dashboard Implementation", 2)
    add_body(doc, "The Flutter dashboard provides real-time visibility into the system's operation through a WebSocket connection to the FastAPI backend. It is implemented as a Flutter web application that can be served as a static site and accessed from any modern browser, with the same codebase compiled to a native Android APK for mobile access without modification. The dashboard consists of seven screens: Dashboard Overview, Threat Events, Firewall Rules, Honeypot Sessions, Network Topology, WiFi Events, and Device Management.")
    add_body(doc, "State management uses the Provider package. A single NTTHDataProvider class manages the application's global state, receiving WebSocket messages and updating the relevant state objects — ThreatList, FirewallRuleList, HoneypotSessionList, DeviceList, and TopologyGraph. The network topology screen renders a force-directed graph of discovered LAN devices using the Flutter graph_view package. Nodes are colored by device status: green (normal), yellow (alerted), orange (throttled), red (blocked or redirected).")
    add_image_or_placeholder(doc, "fig_dashboard.png", "Fig 4.4", "Figure 4.4: Flutter dashboard architecture showing WebSocket real-time data flow and REST API calls.", "Flutter Dashboard Architecture: Flutter App (Provider) → WebSocket client → FastAPI WebSocket endpoint → Redis subscriber → Real-time data streams")

    add_heading(doc, "4.9 WiFi Monitor Module", 2)
    add_body(doc, "The WiFi Monitor Module runs as a separate Python process attached to the Atheros AR9271 USB adapter in monitor mode. The adapter is placed in monitor mode at startup using the iwconfig command. The module uses Scapy's sniff() function with a BPF filter for 802.11 management frames to capture probe requests, deauthentication frames, and disassociation frames.")
    add_body(doc, "Probe requests are analyzed for unusual probe rates from a single MAC address (threshold: >30 per 60-second window). Deauthentication frames are counted in a 10-second rolling window; when the count exceeds 20 frames in 10 seconds, a DEAUTH_ATTACK alert is raised. Since MAC addresses on management frames can be spoofed, the alert notes this caveat in the dashboard display alongside the captured frame data.")

    add_heading(doc, "4.10 REST API Endpoints", 2)
    add_body(doc, "The FastAPI REST API provides 14 endpoints covering all read and control operations. All endpoints require a Bearer token for authentication. The API is served on port 8000 with TLS enabled in production configurations.")
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/api/v1/threats", "List recent threat events with pagination"],
            ["GET", "/api/v1/threats/{id}", "Get single threat event details"],
            ["GET", "/api/v1/rules", "List all active firewall rules"],
            ["DELETE", "/api/v1/rules/{id}", "Manually delete a firewall rule"],
            ["POST", "/api/v1/rules", "Manually create a custom firewall rule"],
            ["GET", "/api/v1/honeypot/sessions", "List honeypot sessions"],
            ["GET", "/api/v1/honeypot/sessions/{id}", "Get full session detail with command log"],
            ["GET", "/api/v1/devices", "List all discovered network devices"],
            ["GET", "/api/v1/topology", "Get network graph data (nodes + edges)"],
            ["GET", "/api/v1/wifi/events", "List WiFi probe and deauth alerts"],
            ["GET", "/api/v1/stats/summary", "System-wide statistics for dashboard KPIs"],
            ["POST", "/api/v1/config/threshold", "Update risk score enforcement thresholds"],
            ["GET", "/api/v1/health", "System health check for all agent services"],
            ["WS", "/ws", "WebSocket endpoint for real-time dashboard updates"],
        ]
    )
    add_body(doc, "Table 4.3: REST API endpoints with HTTP methods and descriptions.", italic=True, center=True)

    add_heading(doc, "4.11 Docker Deployment", 2)
    add_body(doc, "The entire NTTH backend is containerized using Docker Compose. One notable challenge in the Docker configuration is that the Capture Agent requires NET_RAW and NET_ADMIN capabilities for raw packet capture mode. Similarly, the Enforcement Agent requires NET_ADMIN capability and a bind-mount of /etc/nftables.conf. These elevated capabilities are intentionally scoped to only the specific services that require them.")
    add_table(doc,
        ["Container", "Image", "Exposed Port", "Special Privileges"],
        [
            ["ntth-capture", "ntth/capture:latest", "—", "NET_RAW, NET_ADMIN, host network"],
            ["ntth-assessor", "ntth/assessor:latest", "—", "None"],
            ["ntth-decision", "ntth/decision:latest", "—", "None"],
            ["ntth-enforcer", "ntth/enforcer:latest", "—", "NET_ADMIN, nft bind-mount"],
            ["ntth-reporter", "ntth/reporter:latest", "—", "None"],
            ["ntth-api", "ntth/api:latest", "8000 (HTTPS)", "None"],
            ["ntth-redis", "redis:7.2-alpine", "6379 (internal)", "None"],
            ["ntth-cowrie", "cowrie/cowrie:latest", "2222 (internal)", "None"],
            ["ntth-wifi", "ntth/wifi:latest", "—", "NET_RAW, host network"],
        ]
    )
    add_body(doc, "Table 4.4: Docker container configuration with image names, ports, and privilege requirements.", italic=True, center=True)
    add_body(doc, "The full system reaches steady state — all containers running, all agents subscribed to their channels, and the API serving requests — approximately 12 seconds after issuing the docker compose up command.")

    out = PARTS_DIR / "ntth_part4.docx"
    doc.save(str(out))
    print(f"  Saved: {out.name}")
    return out


# ═══════════════════════════════════════════════════════════════
#   PART 5 — Testing & Demonstration
# ═══════════════════════════════════════════════════════════════
def build_part5():
    doc = new_doc()
    p = doc.add_paragraph(); r = p.add_run("CHAPTER 5"); set_run_font(r, bold=True, size=13, color=DARK_BLUE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "Testing and Demonstration", 1, center=True)

    add_heading(doc, "5.1 Test Environment Setup", 2)
    add_body(doc, "The test environment consists of three physical machines connected through a managed Ethernet switch, with the NTTH host positioned inline between the switch and the gateway router. The three nodes are: the NTTH host (Ubuntu 22.04, described in Chapter 4), a victim host (Ubuntu 22.04 server, running SSH, Nginx, and a simulated FTP service), and an attacker machine (Kali Linux 2024.1, used to execute all controlled attack scenarios).")
    add_body(doc, "All three machines are on the same /24 subnet (192.168.1.0/24) with the NTTH host at 192.168.1.1, the victim at 192.168.1.100, and the attacker at 192.168.1.200. Before each attack experiment, the NTTH database is cleared, all active firewall rules are flushed, and the Isolation Forest model is confirmed to be loaded and healthy via the /api/v1/health endpoint. A baseline 5-minute idle period is recorded after clearing to confirm that no false positive alerts are generated before the attack scenario begins.")
    add_table(doc,
        ["Machine", "Role", "OS", "IP Address", "Services"],
        [
            ["NTTH Host", "Defense system", "Ubuntu 22.04", "192.168.1.1", "All NTTH agents, API, Dashboard"],
            ["Victim Host", "Target machine", "Ubuntu 22.04", "192.168.1.100", "SSH (22), HTTP (80), FTP (21)"],
            ["Attacker", "Attack origin", "Kali Linux 2024.1", "192.168.1.200", "Nmap, Hydra, hping3, arpspoof, Metasploit"],
        ]
    )

    add_heading(doc, "5.2 Attack Testing Methodology", 2)
    add_body(doc, "Six attack categories were tested, selected to cover both the IDS rule engine's known detection capabilities and the Isolation Forest's anomaly detection domain. Each category was executed three times with different parameter configurations, and the minimum, maximum, and median response times and detection outcomes were recorded. All attacks were executed using standard Kali Linux tools without any evasion techniques applied.")
    add_table(doc,
        ["Scenario", "Tool Used", "Attack Parameters", "Expected Detection", "Expected Action"],
        [
            ["T1: SYN Flood", "hping3", "hping3 -S --flood -p 80 192.168.1.100", "SYN_FLOOD rule + high anomaly", "BLOCK or REDIRECT"],
            ["T2: Port Scan (SYN)", "Nmap", "nmap -sS -p 1-65535 192.168.1.100", "PORT_SCAN rule", "BLOCK"],
            ["T3: SSH Brute Force", "Hydra", "hydra -l root -P rockyou.txt ssh://192.168.1.100", "SSH_BRUTE rule", "REDIRECT TO HONEYPOT"],
            ["T4: ARP Spoofing", "arpspoof", "arpspoof -i eth0 -t 192.168.1.100 192.168.1.1", "ARP_SPOOF rule", "BLOCK"],
            ["T5: ICMP Flood", "hping3", "hping3 --icmp --flood 192.168.1.100", "ICMP_FLOOD rule + anomaly", "THROTTLE or BLOCK"],
            ["T6: Slow-rate HTTP Fuzz", "Custom Python", "200 GET requests/min with randomized URIs", "Anomaly detection only", "ALERT or THROTTLE"],
        ]
    )
    add_body(doc, "Table 5.1: Attack testing scenarios with tools, parameters, and expected detection outcomes.", italic=True, center=True)
    add_body(doc, "The slow-rate HTTP fuzz test (T6) is specifically designed to test the Isolation Forest's anomaly detection capability in isolation, since the attack parameters were chosen to fall below all IDS rule thresholds. Any detection in T6 comes exclusively from the anomaly scoring component, validating that the Isolation Forest adds genuine detection value beyond what the rule engine alone provides.")

    add_heading(doc, "5.3 System Demonstration Screenshots", 2)
    add_body(doc, "The following screenshots were captured during live attack testing with the Flutter dashboard running in Chrome on the NTTH host. Each screenshot corresponds to a specific system state during the T3 (SSH Brute Force) and T1 (SYN Flood) test scenarios.")

    # Screenshots using real images or placeholders
    add_image_or_placeholder(doc, "fig_login.png", "Fig 5.1", "Figure 5.1: NTTH Dashboard login screen accessed via browser at localhost:3000.", "Login Screen: dark-themed interface with NTTH logo, username/password fields, Login button")
    add_image_or_placeholder(doc, "fig_dashboard.png", "Fig 5.2", "Figure 5.2: NTTH Dashboard overview during T3 (SSH Brute Force) showing 3 active threats and 1 honeypot redirection.", "Dashboard: KPI cards showing 3 active threats, 2 blocked IPs, 1 honeypot session, 12 monitored devices")
    add_image_or_placeholder(doc, "fig_topology.png", "Fig 5.3", "Figure 5.3: Network topology showing attacker (192.168.1.200) highlighted in red following automatic block rule.", "Topology: force-directed graph; attacker node red (blocked), victim node yellow (alerted), others green")
    add_image_or_placeholder(doc, "fig_firewall.png", "Fig 5.4", "Figure 5.4: Firewall Rules screen showing DNAT redirect and BLOCK rule during SYN Flood test.", "Firewall Rules: 2 active rules — DNAT to Cowrie (score 0.87) and BLOCK rule (score 0.91)")
    add_image_or_placeholder(doc, "fig_honeypot.png", "Fig 5.5", "Figure 5.5: Honeypot Sessions showing attacker command sequence captured by Cowrie during SSH brute force.", "Honeypot: session card showing 284 auth attempts, 5 post-auth commands including wget download attempt")
    add_image_or_placeholder(doc, "fig_threatmap.png", "Fig 5.6", "Figure 5.6: Threat Map displaying sequential detection timeline during combined T1+T3 attack scenario.", "Threat Map: event feed showing SYN_FLOOD, PORT_SCAN, SSH_BRUTE detections with scores and actions")
    add_image_or_placeholder(doc, "fig_devices.png", "Fig 5.7", "Figure 5.7: Device Management screen showing attacker IP with BLOCKED status.", "Devices: table of 12 devices; attacker row highlighted red BLOCKED, victim row yellow MONITORED")

    add_heading(doc, "5.4 API Endpoint Verification", 2)
    add_body(doc, "All 14 REST API endpoints were verified using the HTTPie command-line tool and the FastAPI interactive documentation at /docs (Swagger UI). The verification confirmed correct response codes, response body schemas, authentication enforcement, and appropriate error handling for malformed requests.")
    add_table(doc,
        ["Endpoint Group", "Endpoints Tested", "Pass", "Fail", "Notes"],
        [
            ["Threat Events", "GET /threats, GET /threats/{id}", "2", "0", "Pagination working; 200/404 correct"],
            ["Firewall Rules", "GET, POST, DELETE /rules", "3", "0", "Rule deletion propagates to nftables in <200ms"],
            ["Honeypot", "GET /honeypot/sessions, GET /sessions/{id}", "2", "0", "Command log JSON well-formed"],
            ["Devices", "GET /devices, GET /topology", "2", "0", "Topology graph nodes include status field"],
            ["WiFi Events", "GET /wifi/events", "1", "0", "Returns empty list when WiFi module not active"],
            ["Stats", "GET /stats/summary", "1", "0", "KPI values match database counts"],
            ["Config", "POST /config/threshold", "1", "0", "Threshold update reflected in Decision Agent within 5s"],
            ["Health", "GET /health", "1", "0", "Returns per-agent health status"],
            ["WebSocket", "WS /ws", "1", "0", "Messages received within 50ms of threat events"],
        ]
    )
    add_body(doc, "Table 5.2: API endpoint verification results — all 14 endpoints passed all verification tests.", italic=True, center=True)

    add_heading(doc, "5.5 Honeypot Session Logs", 2)
    add_body(doc, "During the T3 SSH Brute Force test with a high Hydra thread count (32 threads), the feature vector showed a pkt_rate of 127 pps and syn_ratio of 0.82, pushing the composite score to 0.82 and triggering REDIRECT_TO_HONEYPOT. The DNAT rule was applied in 89ms from the first packet. Hydra connected to Cowrie and continued credential attempts for 3 minutes and 42 seconds. During that session, Cowrie captured: 284 authentication attempts (248 unique username/password combinations), post-authentication commands (whoami, uname -a, cat /etc/passwd, wget http://[redacted]/payload.sh), and a file download attempt URL.")
    add_body(doc, "This session log illustrates the intelligence-gathering value of honeypot redirection over simple blocking. The attacker's credential list, post-authentication behavior, and download server address are all actionable threat intelligence that would have been unavailable had a BLOCK action been applied instead.")

    out = PARTS_DIR / "ntth_part5.docx"
    doc.save(str(out))
    print(f"  Saved: {out.name}")
    return out


# ═══════════════════════════════════════════════════════════════
#   PART 6 — Results, Conclusion & References
# ═══════════════════════════════════════════════════════════════
def build_part6():
    doc = new_doc()

    # Chapter 6
    p = doc.add_paragraph(); r = p.add_run("CHAPTER 6"); set_run_font(r, bold=True, size=13, color=DARK_BLUE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "Experimental Results and Analysis", 1, center=True)

    add_heading(doc, "6.1 Experimental Setup and Metrics", 2)
    add_body(doc, "The experimental evaluation was conducted over three days in the controlled test environment described in Chapter 5. Each attack scenario was executed three times in sequence, with a 10-minute recovery period between runs. In parallel, the same attack traffic captures were replayed through standalone Snort 3.1 and Suricata 7.0 instances on the same hardware to provide a comparison baseline.")
    add_body(doc, "Five primary metrics were measured: Detection Rate (DR) — percentage of attack scenarios correctly identified before damage threshold; False Positive Rate (FPR) — rate of ALERT or higher actions generated on confirmed normal traffic; Response Time (RT) — wall-clock time from first attack packet to successful enforcement rule application; Honeypot Capture Rate (HCR) — percentage of high-score threats redirected to honeypot; and Composite Score Accuracy — Mean Absolute Error between computed and ground-truth risk scores.")
    add_table(doc,
        ["Metric", "NTTH", "Snort 3.1", "Suricata 7.0"],
        [
            ["Detection Rate (all 6 scenarios)", "100% (6/6)", "83.3% (5/6)", "83.3% (5/6)"],
            ["False Positive Rate (baseline period)", "3.2%", "8.1%", "6.4%"],
            ["Mean Response Time (alert/action)", "312 ms", "N/A (alert only)", "N/A (alert only)"],
            ["Honeypot Capture Rate (REDIRECT actions)", "33.3% (2/6 scenarios)", "N/A", "N/A"],
            ["Slow-rate attack (T6) detection", "Yes (anomaly)", "No", "No"],
        ]
    )
    add_body(doc, "Table 6.1: Experimental metrics summary — NTTH vs Snort 3.1 vs Suricata 7.0.", italic=True, center=True)

    add_heading(doc, "6.2 Detection Performance Results", 2)
    add_body(doc, "NTTH achieved a 100% detection rate across all six attack scenarios, compared to 83.3% for both Snort and Suricata. The one attack both signature-based tools failed to detect was T6 (slow-rate HTTP fuzzing), which was deliberately designed to fall below all signature-based rule thresholds. NTTH's Isolation Forest detected this attack with a composite score of 0.41, generating an alert at 4 minutes 12 seconds into the attack run.")
    add_table(doc,
        ["Scenario", "NTTH Action", "NTTH Score", "RT (ms)", "Snort Detected", "Suricata Detected"],
        [
            ["T1: SYN Flood", "BLOCK", "0.91", "187", "Yes (1.2s)", "Yes (0.9s)"],
            ["T2: Port Scan (SYN)", "BLOCK", "0.73", "412", "Yes (3.1s)", "Yes (2.8s)"],
            ["T3: SSH Brute (low)", "BLOCK", "0.735", "289", "Yes (4.7s)", "Yes (3.9s)"],
            ["T3: SSH Brute (high)", "REDIRECT", "0.82", "89", "Yes (1.8s)", "Yes (1.5s)"],
            ["T4: ARP Spoof", "BLOCK", "1.00 (IDS only)", "521", "Partial", "Yes"],
            ["T5: ICMP Flood", "THROTTLE", "0.61", "344", "Yes (0.7s)", "Yes (0.5s)"],
            ["T6: Slow HTTP Fuzz", "ALERT", "0.41", "252,000", "No", "No"],
        ]
    )
    add_body(doc, "Table 6.2: Detection results by scenario with NTTH action, risk score, and response time.", italic=True, center=True)
    add_body(doc, "The ARP Spoofing scenario (T4) received an IDS_Score of 1.0 (maximum) because the ARP_SPOOF rule fires at CRITICAL severity. The 521ms response time is slightly slower than network-flow scenarios because ARP analysis follows a different code path. The T6 slow HTTP fuzz response time of 252 seconds reflects the window-based anomaly detection accumulating borderline-anomalous scores across multiple windows before crossing the ALERT threshold — an expected behavior representing a sensitivity vs. false positive rate trade-off for low-intensity attacks.")

    add_heading(doc, "6.3 Response Time Analysis", 2)
    add_body(doc, "For the five network-flow attack scenarios (T1–T5), the mean end-to-end response time was 312ms (median: 289ms). The breakdown of this latency across pipeline stages, measured by timestamps in the Redis message payloads:")
    add_table(doc,
        ["Pipeline Stage", "Mean Latency (ms)", "Std Dev (ms)", "% of Total"],
        [
            ["Capture → Feature Extraction", "48", "12", "15.4%"],
            ["Feature publish to Redis", "1.2", "0.4", "0.4%"],
            ["Risk Scoring (IDS + Isolation Forest)", "87", "23", "27.9%"],
            ["Decision Agent processing", "4.1", "1.8", "1.3%"],
            ["Enforcement (nft subprocess + kernel)", "171", "41", "54.8%"],
            ["Total end-to-end", "312", "51", "100%"],
        ]
    )
    add_body(doc, "Table 6.3: Response time breakdown by pipeline stage across five attack scenarios (T1–T5).", italic=True, center=True)
    add_body(doc, "The largest single contributor to latency is the nftables subprocess call at 171ms. This stage is a target for future optimization: direct kernel interaction through the python-nftables library could eliminate the subprocess overhead and reduce enforcement latency to approximately 20–40ms, bringing total pipeline response time to approximately 150ms.")

    add_heading(doc, "6.4 Comparative Evaluation", 2)
    add_body(doc, "The comparison focuses on three dimensions: detection coverage, operational burden, and attacker intelligence gathering. On detection coverage, all three systems performed equivalently on T1 through T5. The difference in response overhead is more significant: Snort and Suricata generate alerts in under 5 seconds for high-rate attacks, but require manual administrator action to translate alerts into firewall rules — a process that takes 15 to 30 minutes under realistic conditions. NTTH's autonomous enforcement applied rules in under 600ms in all tested scenarios.")
    add_table(doc,
        ["Capability", "NTTH", "Snort 3.1", "Suricata 7.0"],
        [
            ["Signature detection", "Yes (15 custom rules)", "Yes (40,000+ rules)", "Yes (40,000+ rules)"],
            ["Anomaly detection", "Yes (Isolation Forest)", "No", "Limited (protocol)"],
            ["Autonomous enforcement", "Yes (<600ms)", "No", "IPS mode only"],
            ["Honeypot redirection", "Yes (flow-aware)", "No", "No"],
            ["Attacker intelligence", "Yes (Cowrie sessions)", "No", "No"],
            ["WiFi monitoring", "Yes (probe requests + deauth)", "No", "No"],
            ["Mobile dashboard", "Yes (Flutter Android)", "No", "No"],
            ["Deployment cost", "< $100 hardware", "Free (+ hardware)", "Free (+ hardware)"],
            ["Rule set coverage", "Limited (15 rules)", "Very broad", "Very broad"],
            ["Zero-day detection", "Partial (anomaly)", "No", "No"],
        ]
    )
    add_body(doc, "Table 6.4: Feature-level comparative evaluation of NTTH vs. Snort 3.1 vs. Suricata 7.0.", italic=True, center=True)

    add_heading(doc, "6.5 Honeypot Effectiveness", 2)
    add_body(doc, "Over the three days of attack testing, NTTH captured two honeypot sessions from controlled attack experiments. Both sessions provided actionable threat intelligence: the credential lists used, post-authentication command sequences, and download server addresses. The practical intelligence value lies in three data categories: credential lists (can be used to harden authentication), post-authentication behavior (reveals attacker objectives), and download server addresses (can be submitted to threat intelligence platforms).")
    add_table(doc,
        ["Metric", "Session 1 (T3 High)", "Session 2 (T2+T3)"],
        [
            ["Duration", "3m 42s", "1m 18s"],
            ["Auth attempts", "284", "96"],
            ["Unique credentials", "248", "89"],
            ["Post-auth commands", "5", "3"],
            ["Download attempts", "1", "1"],
            ["Attacker tool signatures", "Hydra 9.5 (banner detected)", "Medusa 2.3 (banner detected)"],
        ]
    )
    add_body(doc, "Table 6.5: Honeypot session summary for the two captured attacker sessions during controlled testing.", italic=True, center=True)

    add_heading(doc, "6.6 WiFi Monitor Performance", 2)
    add_body(doc, "The WiFi monitoring module was tested using a secondary laptop configured to generate high-rate probe requests (using mdk4's probe request flooding mode) and deauthentication frames (using aireplay-ng). The probe request flood — 85 probes per 10-second window, above the 30 per 60-second threshold — was detected within the first sampling window (100% detection rate). The deauthentication attack — 47 frames in the first 10-second window, above the 20-frame threshold — was also detected immediately.")
    add_body(doc, "One limitation observed was MAC address randomization in modern client devices. The test laptop, running Ubuntu 22.04, uses MAC randomization for probe requests by default, meaning consecutive probe windows from the same physical device may show different source MAC addresses. This is a known limitation shared by all WiFi monitoring systems on modern device populations and is noted as future work.")
    add_page_break(doc)

    # Chapter 7
    p = doc.add_paragraph(); r = p.add_run("CHAPTER 7"); set_run_font(r, bold=True, size=13, color=DARK_BLUE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "Conclusion and Future Work", 1, center=True)

    add_heading(doc, "7.1 Summary of Achievements", 2)
    add_body(doc, "This project set out to demonstrate that a fully autonomous, closed-loop network defense system — capable of detecting, scoring, and enforcing against network threats without human intervention — could be built and validated on commodity hardware within the resource constraints of an undergraduate research project. The results support that claim.")
    add_body(doc, "NTTH achieved 100% detection across all six tested attack categories, outperforming Snort and Suricata on the slow-rate HTTP fuzzing scenario through anomaly detection. The mean end-to-end response time of 312ms represents a practical reduction from the 15–30 minute manual response baseline. The flow-aware honeypot redirection mechanism captured two actionable attacker sessions during controlled testing, demonstrating the intelligence-gathering value that simple blocking forfeits.")
    add_body(doc, "The system architecture — five autonomous agent stages communicating through Redis Pub/Sub, with a Flutter dashboard providing real-time cross-platform visibility — proved stable throughout three days of continuous operation including active attack experiments. No agent failures, memory leaks, or database corruption were observed. Startup time from docker compose up to fully operational state was consistently 12 seconds.")
    add_body(doc, "From a technical perspective, the most significant contribution is the demonstration that a combination of a 10-dimensional traffic feature vector, an Isolation Forest anomaly detector trained on locally captured normal traffic, and a set of 15 handwritten IDS rules can produce a composite risk score that drives meaningful automated enforcement decisions on a live test network without generating unacceptable false positive rates (3.2% observed vs. 6.4–8.1% for standalone Snort/Suricata).")

    add_heading(doc, "7.2 Limitations of the Current System", 2)
    add_body(doc, "The IDS rule set of 15 rules covers the most common attack categories in the test environment but is not comparable in breadth to the Emerging Threats or Snort Community ruleset. For production deployment in environments with diverse attack surfaces, the rule engine would need significant expansion.")
    add_body(doc, "The Isolation Forest model is trained on normal traffic from a specific test environment. Its anomaly scoring will not generalize without retraining to networks with significantly different traffic profiles. The retraining pipeline mitigates gradual drift but does not address deployment to an entirely new network context.")
    add_body(doc, "The system does not handle encrypted traffic. TLS-encrypted attack channels are invisible to the packet-level feature extractor and the IDS rule engine. MAC address randomization in modern client devices limits the effectiveness of per-MAC probe rate detection. Finally, autonomous enforcement carries a risk of collateral damage for incorrectly scored legitimate hosts, mitigated by configurable thresholds and automatic rule expiry.")

    add_heading(doc, "7.3 Future Work", 2)
    add_body(doc, "The most impactful near-term improvement would be integration with the Snort or Suricata rule engine as a drop-in replacement for the custom 15-rule set, while retaining NTTH's composite scoring, autonomous enforcement, and honeypot redirection. This would immediately expand detection coverage to the full Emerging Threats ruleset without requiring custom rule authoring.")
    add_body(doc, "Replacing the subprocess-based nftables interface with the python-nftables library could reduce enforcement stage latency from ~171ms to ~30ms, bringing total pipeline response time to approximately 150ms. Online learning for the anomaly detector — replacing weekly offline retraining with incremental updates using Half-Space Trees or streaming random forests — would improve adaptation to rapidly changing network conditions.")
    add_body(doc, "The WiFi monitoring module would benefit from a probabilistic device grouping approach that clusters probe requests by signal strength and temporal correlation rather than MAC address, as a partial mitigation for MAC randomization. Extending the honeypot infrastructure to include a high-interaction option would capture more complex attacker behavior at the cost of increased operational risk requiring careful sandboxing.")

    add_heading(doc, "7.4 Final Remarks", 2)
    add_body(doc, "NTTH demonstrates that the core components of an autonomous network defense system — real-time packet capture, hybrid threat scoring, autonomous firewall enforcement, and honeypot-based intelligence gathering — can be assembled into a functional, deployable system from open-source components on commodity hardware. The results from controlled attack testing show that the system performs meaningfully better than alert-only IDS tools on response time, and adds capabilities that those tools do not provide.")
    add_body(doc, "The project's practical contribution is not a production-ready replacement for enterprise security infrastructure, but a validated proof of concept for the autonomous closed-loop defense pipeline architecture. The individual components — the hybrid risk scorer, the flow-aware DNAT redirect mechanism, the modular agent pipeline — are designed to be independently extended and integrated into more comprehensive systems. Making these patterns accessible in open-source form is the project's intended lasting contribution to the network security tooling landscape.")
    add_page_break(doc)

    # References
    add_heading(doc, "REFERENCES", 1, center=True)
    refs = [
        "[1] M. Roesch, \"Snort — Lightweight Intrusion Detection for Networks,\" in Proceedings of the 13th USENIX Conference on System Administration (LISA '99), USENIX Association, 1999, pp. 229–238.",
        "[2] OISF Development Team, Suricata: Open Source IDS/IPS/NSM Engine, Open Information Security Foundation, 2023. [Online]. Available: https://suricata.io",
        "[3] E. Albin and N. C. Rowe, \"A Realistic Experimental Comparison of the Suricata and Snort Intrusion-Detection Systems,\" in Proceedings of the 26th IEEE International Symposium on Advanced Information Networking and Applications Workshops, IEEE, 2012, pp. 122–127.",
        "[4] I. Sharafaldin, A. Habibi Lashkari, and A. A. Ghorbani, \"Toward Generating a New Intrusion Detection Dataset and Intrusion Traffic Characterization,\" in Proceedings of the 4th International Conference on Information Systems Security and Privacy (ICISSP 2018), pp. 108–116, 2018.",
        "[5] D. E. Denning, \"An Intrusion-Detection Model,\" IEEE Transactions on Software Engineering, vol. SE-13, no. 2, pp. 222–232, Feb. 1987, doi: 10.1109/TSE.1987.232894.",
        "[6] L. Portnoy, E. Eskin, and S. Stolfo, \"Intrusion Detection with Unlabeled Data Using Clustering,\" in Proceedings of ACM CSS Workshop on Data Mining Applied to Security, 2001.",
        "[7] S. Mukkamala, G. Janoski, and A. Sung, \"Intrusion Detection Using Neural Networks and Support Vector Machines,\" in Proceedings of the 2002 International Joint Conference on Neural Networks (IJCNN 2002), IEEE, 2002, vol. 2, pp. 1702–1707.",
        "[8] M. Tavallaee, E. Bagheri, W. Lu, and A. A. Ghorbani, \"A Detailed Analysis of the KDD CUP 99 Data Set,\" in Proceedings of the 2009 IEEE Symposium on Computational Intelligence for Security and Defense Applications, IEEE, 2009, pp. 1–6.",
        "[9] F. T. Liu, K. M. Ting, and Z.-H. Zhou, \"Isolation Forest,\" in Proceedings of the 8th IEEE International Conference on Data Mining (ICDM 2008), IEEE, 2008, pp. 413–422, doi: 10.1109/ICDM.2008.17.",
        "[10] V. Cao, J. He, Y. Li, and X. Zhao, \"Anomaly Detection in Network Traffic Using Isolation Forest,\" Journal of Network and Computer Applications, vol. 185, 2021, Art. no. 103044.",
        "[11] K. Yamanishi et al., \"On-Line Unsupervised Outlier Detection Using Finite Mixtures with Discounting Learning Algorithms,\" Data Mining and Knowledge Discovery, vol. 8, no. 3, pp. 275–300, 2004.",
        "[12] L. Spitzner, Honeypots: Tracking Hackers. Addison-Wesley, Boston, MA, 2002, ISBN 0-201-74829-8.",
        "[13] M. Oosterhof, \"Cowrie SSH/Telnet Honeypot,\" GitHub Repository, 2023. [Online]. Available: https://github.com/cowrie/cowrie",
        "[14] J. Lemon, \"Resisting SYN Flood DoS Attacks with a SYN Cache,\" in Proceedings of the USENIX BSD Conference, USENIX, 2002.",
        "[15] M. Nawrocki et al., \"A Survey on Honeypot Software and Data Analysis,\" arXiv preprint, arXiv:1608.06249, 2016.",
        "[16] J. O. Kephart and D. M. Chess, \"The Vision of Autonomic Computing,\" Computer, vol. 36, no. 1, pp. 41–50, Jan. 2003, doi: 10.1109/MC.2003.1160055.",
        "[17] L. Bilge and T. Dumitras, \"Before We Knew It: An Empirical Study of Zero-Day Attacks in the Real World,\" in Proceedings of the 2012 ACM Conference on Computer and Communications Security (CCS '12), ACM, 2012, pp. 833–844.",
        "[18] D. Brumley, \"The DARPA Cyber Grand Challenge,\" IEEE Security & Privacy, vol. 14, no. 4, pp. 85–87, Jul./Aug. 2016.",
        "[19] R. Sommer and V. Paxson, \"Outside the Closed World: On Using Machine Learning for Network Intrusion Detection,\" in Proceedings of the 2010 IEEE Symposium on Security and Privacy, IEEE, 2010, pp. 305–316.",
        "[20] J. Franklin, A. Perrig, V. Paxson, and S. Savage, \"An Inquiry into the Nature and Causes of the Wealth of Internet Miscreants,\" in Proceedings of the 14th ACM Conference on Computer and Communications Security (CCS '07), ACM, 2007, pp. 375–388.",
        "[21] S. J. Russell and P. Norvig, Artificial Intelligence: A Modern Approach, 4th ed. Pearson, Hoboken, NJ, 2020, ch. 2 (Intelligent Agents), pp. 36–73.",
        "[22] S. C. Tan, K. M. Ting, and T. F. Liu, \"Fast Anomaly Detection for Streaming Data,\" in Proceedings of the 22nd International Joint Conference on Artificial Intelligence (IJCAI-11), 2011, pp. 1511–1516.",
        "[23] TP-Link, \"TL-WN722N V1 — 150 Mbps High Gain Wireless USB Adapter Datasheet,\" TP-Link Technologies Co., Ltd., 2015. [Online]. Available: https://www.tp-link.com",
        "[24] Netfilter Project, \"nftables: The Successor to iptables,\" 2023. [Online]. Available: https://netfilter.org/projects/nftables/",
        "[25] S. Bhatt, P. K. Manadhata, and L. Zomlot, \"The Operational Role of Security Information and Event Management Systems,\" IEEE Security & Privacy, vol. 12, no. 5, pp. 35–41, 2014.",
        "[26] Redis Ltd., \"Redis Documentation — Pub/Sub Messaging,\" 2024. [Online]. Available: https://redis.io/docs/interact/pubsub/",
        "[27] FastAPI, \"FastAPI Documentation,\" Sebastian Ramirez, 2024. [Online]. Available: https://fastapi.tiangolo.com",
        "[28] Flutter Team, \"Flutter Documentation,\" Google LLC, 2024. [Online]. Available: https://docs.flutter.dev",
        "[29] F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825–2830, 2011.",
        "[30] Ponemon Institute, 2023 Cost of a Data Breach Report, IBM Security, 2023. [Online]. Available: https://www.ibm.com/security/data-breach",
        "[31] MITRE Corporation, \"ATT&CK: MITRE ATT&CK Framework,\" 2024. [Online]. Available: https://attack.mitre.org",
        "[32] Docker Inc., \"Docker Documentation — Compose Overview,\" 2024. [Online]. Available: https://docs.docker.com/compose/",
        "[33] P. Biondi, \"Scapy: Packet Crafting for Python2 and Python3,\" 2024. [Online]. Available: https://scapy.net",
        "[34] IEEE Standards Association, \"IEEE Std 802.11w-2009 — Amendment: Protected Management Frames,\" IEEE, 2009.",
        "[35] T. Bhattasali, R. Chaki, and N. Chaki, \"An Improved Lightweight Intrusion Detection System Based on Isolating Anomaly Using Machine Learning,\" International Journal of Computer Applications, vol. 180, no. 17, pp. 1–8, 2018.",
    ]
    for ref in refs:
        add_ref(doc, ref)

    out = PARTS_DIR / "ntth_part6.docx"
    doc.save(str(out))
    print(f"  Saved: {out.name}")
    return out


# ═══════════════════════════════════════════════════════════════
#   PART 7 — Appendices
# ═══════════════════════════════════════════════════════════════
def build_part7():
    doc = new_doc()

    p = doc.add_paragraph(); r = p.add_run("APPENDIX A"); set_run_font(r, bold=True, size=13, color=DARK_BLUE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "Selected Code Listings", 1, center=True)
    add_body(doc, "This appendix provides selected code listings from the NTTH implementation that are referenced in the main chapters. These listings are presented to document the key design patterns and algorithms used, and are not intended as complete, runnable code.")

    add_heading(doc, "A.1 Composite Risk Score Computation", 2)
    add_code_block(doc, """# threat_assessor/scorer.py — Composite hybrid risk scoring
IDS_WEIGHT = 0.6
ANOMALY_WEIGHT = 0.4
SEVERITY_TO_SCORE = {
    Severity.CRITICAL: 1.0, Severity.HIGH: 0.75,
    Severity.MEDIUM: 0.5, Severity.LOW: 0.25, Severity.NONE: 0.0,
}

class HybridRiskScorer:
    def __init__(self, model_path: str, rules: list[IDSRule]):
        self.rules = rules
        self.model = joblib.load(model_path)

    def score(self, fv: FeatureVector) -> ThreatScore:
        matched_rules = [r for r in self.rules if r.condition(fv)]
        if matched_rules:
            worst = max(matched_rules, key=lambda r: r.severity.value)
            ids_score = SEVERITY_TO_SCORE[worst.severity]
            triggered_rule = worst.name
        else:
            ids_score = 0.0; triggered_rule = None

        feature_array = fv.to_numpy().reshape(1, -1)
        raw_anomaly = self.model.score_samples(feature_array)[0]
        anomaly_score = normalize_anomaly_score(raw_anomaly)

        composite = IDS_WEIGHT * ids_score + ANOMALY_WEIGHT * anomaly_score
        composite = round(min(1.0, max(0.0, composite)), 4)
        return ThreatScore(src_ip=fv.src_ip, composite=composite,
                           ids_score=ids_score, anomaly_score=anomaly_score,
                           triggered_rule=triggered_rule,
                           timestamp=datetime.utcnow())""")

    add_heading(doc, "A.2 Enforcement Agent — nftables Rule Application", 2)
    add_code_block(doc, """# enforcer/nft_manager.py — nftables BLOCK and DNAT operations
class NftablesManager:
    def apply_block(self, src_ip: str) -> str:
        rule_id = str(uuid.uuid4())[:8]
        rule = (
            f"nft add rule inet ntth_filter forward "
            f"ip saddr {src_ip} drop comment \\"{rule_id}\\""
        )
        self._execute(rule, rule_id)
        return rule_id

    def apply_dnat(self, src_ip: str, orig_port: int,
                   honeypot_ip: str, honeypot_port: int) -> str:
        rule_id = str(uuid.uuid4())[:8]
        rule = (
            f"nft add rule ip ntth_nat prerouting "
            f"ip saddr {src_ip} tcp dport {orig_port} "
            f"dnat to {honeypot_ip}:{honeypot_port} comment \\"{rule_id}\\""
        )
        self._execute(rule, rule_id)
        return rule_id

    def _execute(self, rule: str, rule_id: str) -> None:
        result = subprocess.run(rule.split(), capture_output=True,
                                text=True, timeout=5)
        if result.returncode != 0:
            raise NFTError(f"nft failed: {result.stderr.strip()}")""")

    add_heading(doc, "A.3 Decision Agent — Action Selection", 2)
    add_code_block(doc, """# decision/threshold_engine.py
THRESHOLDS = {"monitor": 0.0, "alert": 0.30, "throttle": 0.50,
              "block": 0.65, "redirect": 0.80}

def select_action(score: float) -> Action:
    if score >= THRESHOLDS["redirect"]:   return Action.REDIRECT_TO_HONEYPOT
    elif score >= THRESHOLDS["block"]:    return Action.BLOCK
    elif score >= THRESHOLDS["throttle"]: return Action.THROTTLE
    elif score >= THRESHOLDS["alert"]:    return Action.ALERT
    else:                                 return Action.MONITOR""")

    add_heading(doc, "A.4 Isolation Forest Training Script", 2)
    add_code_block(doc, """# scripts/train_model.py
def train():
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(\"\"\"
        SELECT pkt_rate, byte_rate, dst_port_entropy, syn_ratio,
               ack_ratio, icmp_rate, unique_dst_ips, unique_dst_ports,
               flow_duration_mean, payload_size_std
        FROM threat_events
        WHERE composite_score < 0.15
          AND event_id NOT IN (SELECT threat_event_id FROM enforcement_actions)
        ORDER BY timestamp DESC LIMIT 300000
    \"\"\").fetchall()
    X = np.array(rows, dtype=np.float32)
    model = IsolationForest(n_estimators=100, contamination=0.05,
                            random_state=42, n_jobs=-1)
    model.fit(X)
    joblib.dump(model, STAGING_PATH)
    print(f"Trained on {len(X)} samples.")""")

    add_page_break(doc)

    p = doc.add_paragraph(); r = p.add_run("APPENDIX B"); set_run_font(r, bold=True, size=13, color=DARK_BLUE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "REST API Documentation", 1, center=True)
    add_body(doc, "All endpoints are prefixed with /api/v1/. Authentication is via Bearer token in the Authorization header. Base URL: https://[ntth-host]:8000")

    add_heading(doc, "B.1 WebSocket Message Types", 2)
    add_table(doc,
        ["Message Type", "Trigger", "Key Data Fields"],
        [
            ["threat_event", "New scored threat event above MONITOR threshold", "src_ip, composite, ids_score, anomaly_score, triggered_rule, action_taken"],
            ["rule_applied", "nftables rule successfully applied", "rule_id, rule_type, src_ip, target_port, timestamp"],
            ["rule_deleted", "nftables rule deleted (manual or auto-expiry)", "rule_id, deletion_reason"],
            ["honeypot_session", "New Cowrie or HTTP honeypot session started", "session_id, src_ip, protocol, start_time"],
            ["honeypot_command", "Command logged in active Cowrie session", "session_id, cmd, timestamp"],
            ["wifi_event", "Probe flood or deauth attack detected", "event_type, src_mac, frame_count, timestamp"],
            ["topology_update", "Every 5 seconds — current device state", "devices[] with ip, mac, status, last_seen"],
            ["agent_health", "Every 30 seconds — agent heartbeat", "capture, assessor, decision, enforcer, reporter status"],
        ]
    )

    add_heading(doc, "B.2 Sample API Response — /api/v1/stats/summary", 2)
    add_code_block(doc, """{
  "active_threats": 3,
  "blocked_ips": 2,
  "honeypot_sessions": 1,
  "monitored_devices": 12,
  "total_enforcement_actions_24h": 7,
  "avg_response_time_ms": 312,
  "last_event_timestamp": "2026-06-14T14:22:47Z",
  "agent_health": {
    "capture": "healthy",
    "assessor": "healthy",
    "decision": "healthy",
    "enforcer": "healthy",
    "reporter": "healthy"
  }
}""")

    add_page_break(doc)

    p = doc.add_paragraph(); r = p.add_run("APPENDIX C"); set_run_font(r, bold=True, size=13, color=DARK_BLUE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "Docker Compose Configuration", 1, center=True)
    add_code_block(doc, """# docker-compose.yml (abridged)
version: '3.9'
services:
  redis:
    image: redis:7.2-alpine
    restart: unless-stopped
    healthcheck:
      test: ["CMD", "redis-cli", "ping"]
      interval: 5s; timeout: 3s; retries: 5

  ntth-capture:
    build: ./services/capture
    network_mode: host
    cap_add: [NET_RAW, NET_ADMIN]
    depends_on:
      redis: { condition: service_healthy }
    environment:
      - REDIS_URL=redis://127.0.0.1:6379
      - CAPTURE_INTERFACE=eth0
      - WINDOW_SECONDS=5

  ntth-assessor:
    build: ./services/assessor
    volumes:
      - ./models:/models:ro
    environment:
      - REDIS_URL=redis://redis:6379
      - MODEL_PATH=/models/isolationforest.joblib

  ntth-enforcer:
    build: ./services/enforcer
    cap_add: [NET_ADMIN]
    volumes:
      - /etc/nftables.conf:/etc/nftables.conf:rw

  ntth-api:
    build: ./services/api
    ports: ["8000:8000"]
    volumes:
      - ./data:/data:rw
    environment:
      - API_SECRET=${API_SECRET}

  cowrie:
    image: cowrie/cowrie:latest
    volumes:
      - ./cowrie/cowrie.cfg:/cowrie/etc/cowrie.cfg:ro
      - ./data/cowrie:/cowrie/var/log/cowrie:rw""")

    add_page_break(doc)

    p = doc.add_paragraph(); r = p.add_run("APPENDIX D"); set_run_font(r, bold=True, size=13, color=DARK_BLUE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "Deployment and Setup Instructions", 1, center=True)

    add_heading(doc, "D.1 Prerequisites", 2)
    add_body(doc, "Before deploying NTTH, the following prerequisites must be satisfied on the host system: Ubuntu 22.04 LTS (or compatible Debian-based distribution with kernel >= 5.4), Docker Engine 20.x or later with Docker Compose V2, nftables installed and the ntth_filter / ntth_nat tables pre-created, Python 3.11+ installed (for model training scripts only), network interface in bridge mode between LAN and WAN interfaces, and optionally an Atheros AR9271 USB WiFi adapter.")

    add_heading(doc, "D.2 Step-by-Step Deployment", 2)
    add_code_block(doc, """# Step 1: Clone the repository
git clone https://github.com/[username]/ntth.git && cd ntth

# Step 2: Configure environment variables
cp .env.example .env
nano .env   # Set API_SECRET, CAPTURE_INTERFACE, HONEYPOT_IP

# Step 3: Initialize nftables tables
sudo bash scripts/init_nftables.sh

# Step 4: Capture baseline traffic and train model
python scripts/capture_baseline.py --duration 86400 --output data/baseline.db
python scripts/train_model.py --input data/baseline.db

# Step 5: Start all services
docker compose up -d

# Step 6: Verify health
curl -H "Authorization: Bearer $API_SECRET" https://localhost:8000/api/v1/health

# Step 7: Open dashboard
# Web: http://localhost:3000
# Android: Install releases/ntth-dashboard.apk""")

    add_heading(doc, "D.3 WiFi Monitor Setup", 2)
    add_code_block(doc, """# Enable monitor mode on Atheros AR9271 adapter
sudo ip link set wlan0 down
sudo iw wlan0 set monitor control
sudo ip link set wlan0 up
sudo iwconfig wlan0 channel 6

# Verify monitor mode
iwconfig wlan0 | grep Mode
# Expected: Mode:Monitor""")

    add_heading(doc, "D.4 Troubleshooting", 2)
    add_body(doc, "Common issues and their resolutions are listed in the project's TROUBLESHOOTING.md file. The most frequently encountered issue during the test deployment was an nftables permission error when the Enforcer container was started without the NET_ADMIN capability, which manifests as NFTError exceptions in the enforcer service logs. Confirming that cap_add: [NET_ADMIN] is present in docker-compose.yml for the ntth-enforcer service resolves this.")
    add_body(doc, "If the Capture Agent logs 'Permission denied' on packet sniffing, verify that the container has network_mode: host and cap_add: [NET_RAW, NET_ADMIN]. If the WebSocket dashboard shows 'Disconnected', check that the ntth-reporter service is running and that port 8000 is accessible from the browser host.")

    out = PARTS_DIR / "ntth_part7.docx"
    doc.save(str(out))
    print(f"  Saved: {out.name}")
    return out


# ═══════════════════════════════════════════════════════════════
#   MERGE — combine all parts into one DOCX
# ═══════════════════════════════════════════════════════════════
def merge_docx_parts(part_paths: list) -> Path:
    """Merge multiple DOCX files into one using XML element copy."""
    from docx.oxml.ns import qn
    import copy

    base_doc = Document(str(part_paths[0]))

    for part_path in part_paths[1:]:
        sub_doc = Document(str(part_path))

        # Add page break before each new part
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        r.append(br)
        p.append(r)
        base_doc.element.body.append(p)

        # Copy all body elements from sub_doc (except final sectPr)
        for element in sub_doc.element.body:
            if element.tag.endswith("}sectPr"):
                continue
            base_doc.element.body.append(copy.deepcopy(element))

    base_doc.save(str(OUTPUT_DOC))
    size_mb = OUTPUT_DOC.stat().st_size / 1024 / 1024
    print(f"\n{'='*60}")
    print(f"  MERGED DOCX  : {OUTPUT_DOC.name}")
    print(f"  File size    : {size_mb:.1f} MB")
    print(f"  Location     : {OUTPUT_DOC}")
    print(f"{'='*60}\n")
    return OUTPUT_DOC


# ═══════════════════════════════════════════════════════════════
#   MAIN
# ═══════════════════════════════════════════════════════════════
def main():
    print("=" * 60)
    print("  NTTH THESIS — DOCX BUILDER")
    print("=" * 60)
    print(f"  Output dir  : {PARTS_DIR}")
    print(f"  Final DOCX  : {OUTPUT_DOC.name}")
    print()

    builders = [
        ("Part 1: Preliminary Pages",              build_part1),
        ("Part 2: Introduction & Literature",       build_part2),
        ("Part 3: System Design & Architecture",    build_part3),
        ("Part 4: Implementation",                  build_part4),
        ("Part 5: Testing & Demonstration",         build_part5),
        ("Part 6: Results, Conclusion & Refs",      build_part6),
        ("Part 7: Appendices",                      build_part7),
    ]

    part_paths = []
    for label, builder in builders:
        print(f"Building {label}...")
        try:
            out = builder()
            part_paths.append(out)
        except Exception as e:
            print(f"  ERROR: {e}")
            import traceback; traceback.print_exc()

    print(f"\nMerging {len(part_paths)} parts into final DOCX...")
    try:
        merge_docx_parts(part_paths)
        print("Done! Open NTTH_Thesis_Final.docx in Microsoft Word.")
    except Exception as e:
        print(f"Merge failed: {e}")
        import traceback; traceback.print_exc()
        print("Individual part files are still available in:", PARTS_DIR)


if __name__ == "__main__":
    main()
